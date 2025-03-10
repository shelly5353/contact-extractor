from flask import Flask, render_template, request, send_file, jsonify
from werkzeug.utils import secure_filename
import os
import tempfile
import logging
import pandas as pd
from datetime import datetime
import docx
import re
import socket
import PyPDF2
import pdfplumber
import phonenumbers
from validate_email import validate_email
import magic
import json

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max-limit
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['DOWNLOAD_FOLDER'] = 'downloads'

# הגדרת לוגר
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# וודא שתיקיות קיימות
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['DOWNLOAD_FOLDER'], exist_ok=True)

ALLOWED_EXTENSIONS = {'xlsx', 'xls', 'doc', 'docx', 'pdf', 'csv'}

# תבניות חיפוש
PATTERNS = {
    'email': r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}',
    'phone': [
        r'(?:\+972|972|05\d|\+972-\d{2}|0\d{1,2}[-\s]?)(?:\d[-\s]?){7,8}',  # ישראלי
        r'\b0\d{1,2}[-\s]?\d{3}[-\s]?\d{4}\b',  # פורמט ישראלי נוסף
        r'\b\d{2,3}[-\s]?\d{3}[-\s]?\d{4}\b',  # מספר מקומי
        r'\b\d{10}\b',  # רצף של 10 ספרות
    ],
    'name': [
        r'[א-ת]+(?: [א-ת]+){1,3}',  # שם בעברית (2-4 מילים)
        r'[A-Z][a-z]+(?: [A-Z][a-z]+){1,3}',  # שם באנגלית (2-4 מילים)
        r'[א-ת]+(?:[-\s][א-ת]+)*',  # שם בעברית עם מקף
        r'[A-Z][a-z]+(?:[-\s][A-Z][a-z]+)*',  # שם באנגלית עם מקף
    ],
    'address': [
        r'(?:[א-ת]+\s)+(?:רחוב|רח\'|שדרות|שד\'|דרך)\s[א-ת\s\d]+(?:\s\d+)?',  # כתובת בעברית
        r'\d+\s(?:[א-ת]+\s)+(?:רחוב|רח\'|שדרות|שד\'|דרך)',  # מספר בית ורחוב
        r'(?:רחוב|רח\'|שדרות|שד\'|דרך)\s[א-ת\s\d]+(?:\s\d+)?',  # כתובת בעברית ללא עיר
    ]
}

def allowed_file(filename):
    """בודק אם הקובץ מורשה"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    """דף הבית"""
    try:
        logger.info("מישהו ניגש לדף הבית")
        return render_template('index.html')
    except Exception as e:
        logger.error(f"שגיאה בטעינת דף הבית: {str(e)}")
        return f"שגיאה בטעינת הדף: {str(e)}", 500

@app.route('/extract', methods=['POST'])
def extract_contacts():
    """חילוץ אנשי קשר מקבצים"""
    if 'files' not in request.files:
        return jsonify({'success': False, 'error': 'לא נבחרו קבצים'})
    
    files = request.files.getlist('files')
    all_contacts = []
    
    try:
        for file in files:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)
                
                logger.info(f"מעבד קובץ: {filename}")
                
                # זיהוי סוג הקובץ
                file_type = magic.from_file(filepath, mime=True)
                contacts = []
                
                if 'pdf' in file_type:
                    contacts = extract_from_pdf(filepath, filename)
                elif 'word' in file_type or 'officedocument' in file_type:
                    contacts = extract_from_word(filepath, filename)
                elif 'excel' in file_type or 'spreadsheet' in file_type:
                    contacts = extract_from_excel(filepath, filename)
                
                all_contacts.extend(contacts)
                os.remove(filepath)
                logger.info(f"קובץ נמחק: {filename}")
        
        if not all_contacts:
            return jsonify({'success': False, 'error': 'לא נמצאו אנשי קשר בקבצים'})
        
        unique_contacts = remove_duplicates(all_contacts)
        output_file = create_excel_output(unique_contacts)
        
        return jsonify({
            'success': True,
            'contacts': unique_contacts[:5],
            'total': len(unique_contacts),
            'download_url': f'/download/{os.path.basename(output_file)}'
        })
        
    except Exception as e:
        logger.error(f"שגיאה בעיבוד הקבצים: {str(e)}")
        return jsonify({'success': False, 'error': str(e)})

def extract_text_from_pdf(filepath):
    """חילוץ טקסט מקובץ PDF"""
    text = ""
    try:
        # ניסיון ראשון עם pdfplumber
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
    except Exception as e:
        logger.warning(f"שגיאה בחילוץ טקסט עם pdfplumber: {str(e)}")
        try:
            # ניסיון שני עם PyPDF2
            with open(filepath, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() or ""
        except Exception as e:
            logger.error(f"שגיאה בחילוץ טקסט עם PyPDF2: {str(e)}")
    
    return text

def extract_from_pdf(filepath, source_filename):
    """חילוץ אנשי קשר מקובץ PDF"""
    text = extract_text_from_pdf(filepath)
    return extract_contacts_from_text(text, source_filename)

def extract_from_word(filepath, source_filename):
    """חילוץ אנשי קשר מקובץ Word"""
    try:
        doc = docx.Document(filepath)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        
        # חילוץ טקסט מטבלאות
        for table in doc.tables:
            for row in table.rows:
                text += '\n' + ' '.join([cell.text for cell in row.cells])
        
        return extract_contacts_from_text(text, source_filename)
    except Exception as e:
        logger.error(f"שגיאה בחילוץ מקובץ Word: {str(e)}")
        return []

def extract_from_excel(filepath, source_filename):
    """חילוץ אנשי קשר מקובץ Excel"""
    contacts = []
    try:
        # קריאת שמות הגליונות
        xlsx = pd.ExcelFile(filepath)
        sheet_names = xlsx.sheet_names
        
        for sheet_name in sheet_names:
            logger.info(f"מעבד גליון: {sheet_name} מתוך הקובץ {source_filename}")
            
            try:
                # קריאת הגליון
                df = pd.read_excel(filepath, sheet_name=sheet_name)
                
                # דילוג על גליונות ריקים
                if df.empty:
                    logger.info(f"גליון ריק: {sheet_name}")
                    continue
                
                # המרת כל העמודות למחרוזות
                df = df.astype(str)
                
                # מיפוי עמודות ידני לפי שמות מוכרים
                column_mapping = {}
                for col in df.columns:
                    col_lower = str(col).lower().strip()
                    if col in ['פרטי', 'שם', 'שם מלא', 'שם פרטי']:
                        column_mapping['name'] = col
                    elif col in ['טלפון', 'נייד', 'טלפון נייד']:
                        column_mapping['phone'] = col
                    elif col in ['אי-מייל', 'מייל', 'אימייל', 'דואר אלקטרוני']:
                        column_mapping['email'] = col
                    elif col in ['עיר', 'כתובת', 'רחוב', 'מען']:
                        column_mapping['address'] = col
                
                logger.info(f"נמצאו העמודות הבאות בגליון {sheet_name}: {column_mapping}")
                
                # עיבוד כל שורה
                for idx, row in df.iterrows():
                    try:
                        # חילוץ נתונים
                        name = str(row[column_mapping.get('name', '')]).strip() if 'name' in column_mapping else ''
                        phone = str(row[column_mapping.get('phone', '')]).strip() if 'phone' in column_mapping else ''
                        email = str(row[column_mapping.get('email', '')]).strip() if 'email' in column_mapping else ''
                        address = str(row[column_mapping.get('address', '')]).strip() if 'address' in column_mapping else ''
                        
                        # ניקוי הטלפון
                        phone = clean_phone(phone)
                        
                        # יצירת איש קשר
                        contact = {
                            'name': name,
                            'phone': phone,
                            'email': email,
                            'address': address,
                            'source': f"{source_filename} - {sheet_name}"
                        }
                        
                        if is_valid_contact(contact):
                            contacts.append(contact)
                            if len(contacts) % 100 == 0:
                                logger.info(f"נמצאו {len(contacts)} אנשי קשר עד כה בקובץ {source_filename}")
                    
                    except Exception as row_error:
                        logger.warning(f"שגיאה בעיבוד שורה {idx} בגליון {sheet_name}: {str(row_error)}")
                        continue
            
            except Exception as sheet_error:
                logger.error(f"שגיאה בעיבוד גליון {sheet_name}: {str(sheet_error)}")
                continue
    
    except Exception as e:
        logger.error(f"שגיאה בחילוץ מקובץ Excel {source_filename}: {str(e)}")
        logger.exception(e)
    
    logger.info(f"נמצאו {len(contacts)} אנשי קשר בקובץ {source_filename}")
    return contacts

def extract_contacts_from_text(text, source_filename):
    """חילוץ אנשי קשר מטקסט"""
    contacts = []
    
    # ניקוי הטקסט
    text = text.replace('\n', ' ').replace('\r', ' ')
    while '  ' in text:
        text = text.replace('  ', ' ')
    
    # חיפוש כל השמות בטקסט
    names = []
    for pattern in PATTERNS['name']:
        names.extend(re.finditer(pattern, text))
    
    # מיון השמות לפי מיקום בטקסט
    names = sorted(names, key=lambda x: x.start())
    
    # עבור כל שם, חיפוש פרטי קשר בסביבתו
    for i, name_match in enumerate(names):
        name = name_match.group()
        
        # קביעת טווח החיפוש - עד השם הבא או 200 תווים
        start_pos = max(0, name_match.start() - 200)
        end_pos = min(len(text), name_match.end() + 200)
        if i < len(names) - 1:
            end_pos = min(end_pos, names[i + 1].start())
        
        surrounding_text = text[start_pos:end_pos]
        
        # חיפוש אימייל
        email = ''
        email_matches = re.finditer(PATTERNS['email'], surrounding_text)
        for match in email_matches:
            potential_email = match.group().lower()
            if validate_email(potential_email):
                email = potential_email
                break
        
        # חיפוש טלפון
        phone = ''
        for pattern in PATTERNS['phone']:
            phone_matches = re.finditer(pattern, surrounding_text)
            for match in phone_matches:
                cleaned_phone = clean_phone(match.group())
                if is_valid_phone(cleaned_phone):
                    phone = cleaned_phone
                    break
            if phone:
                break
        
        # חיפוש כתובת
        address = ''
        for pattern in PATTERNS['address']:
            address_matches = re.finditer(pattern, surrounding_text)
            for match in address_matches:
                address = match.group().strip()
                break
            if address:
                break
        
        # ניקוי השם
        name = name.strip()
        if len(name) < 2 or name.lower() in ['nan', 'none', 'null']:
            continue
        
        contact = {
            'name': name,
            'phone': phone,
            'email': email,
            'address': address,
            'source': source_filename
        }
        
        if is_valid_contact(contact):
            contacts.append(contact)
    
    return contacts

def clean_phone(phone):
    """ניקוי מספר טלפון"""
    if not phone or str(phone).lower() in ['nan', 'none', 'null', '', 'חסר', '-']:
        return ''
    
    # הסרת רווחים מיותרים
    phone = str(phone).strip()
    
    # אם המספר כבר בפורמט תקין (XXX-XXXXXXX), נחזיר אותו כמו שהוא
    if re.match(r'^\d{3}-\d{7}$', phone) or re.match(r'^\d{2}-\d{7}$', phone):
        return phone
    
    # הסרת כל התווים שאינם ספרות
    phone = re.sub(r'[^\d]', '', phone)
    
    # טיפול במספרים בינלאומיים
    if phone.startswith('972'):
        phone = '0' + phone[3:]
    
    # אם המספר לא מתחיל ב-0, נוסיף
    if not phone.startswith('0'):
        phone = '0' + phone
    
    # בדיקת אורך
    if len(phone) < 9 or len(phone) > 10:
        return ''
    
    # בדיקה שזה מספר טלפון ישראלי תקין
    if not re.match(r'^0(([23489]\d{7})|([57]\d{8}))$', phone):
        return ''
    
    # החזרת המספר בפורמט עם מקף
    if len(phone) == 10:  # נייד
        return f"{phone[:3]}-{phone[3:]}"
    else:  # קווי
        return f"{phone[:2]}-{phone[2:]}"

def is_valid_phone(phone):
    """בדיקת תקינות מספר טלפון"""
    if not phone:
        return False
    
    # בדיקה בסיסית למספר ישראלי
    if re.match(r'^0\d{8,9}$', phone):
        return True
    
    try:
        # בדיקה מתקדמת עם ספריית phonenumbers
        parsed_number = phonenumbers.parse(phone, "IL")
        return phonenumbers.is_valid_number(parsed_number)
    except:
        return False

def is_valid_contact(contact):
    """בדיקה אם איש הקשר תקין"""
    # ניקוי השם
    name = contact['name'].strip()
    
    # בדיקות שם
    if not name or name.lower() in ['nan', 'none', 'null', '', 'חסר', '-']:
        return False
    
    # בדיקת אורך מינימלי לשם
    if len(name) < 2:
        return False
    
    # בדיקת טלפון
    phone = contact['phone']
    has_valid_phone = bool(phone and phone != '-' and len(phone) >= 9)
    
    # בדיקת אימייל
    email = contact.get('email', '').strip().lower()
    has_valid_email = bool(
        email and 
        '@' in email and 
        '.' in email.split('@')[1] and 
        email not in ['nan', 'none', 'null', '', 'חסר', '-']
    )
    
    # מספיק שיש או טלפון או אימייל תקין
    return has_valid_phone or has_valid_email

def remove_duplicates(contacts):
    """הסרת כפילויות מרשימת אנשי הקשר"""
    unique_contacts = []
    seen = set()
    
    for contact in contacts:
        # יצירת מפתח ייחודי מהשדות החשובים
        key_parts = [
            contact['name'].lower().strip(),
            contact['phone'].strip() if contact['phone'] else '',
            contact['email'].lower().strip() if contact['email'] else ''
        ]
        key = tuple(key_parts)
        
        if key not in seen:
            seen.add(key)
            unique_contacts.append(contact)
    
    return unique_contacts

def create_excel_output(contacts):
    """יצירת קובץ Excel מעוצב עם אנשי הקשר"""
    filename = f'contacts_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
    output_path = os.path.join(app.config['DOWNLOAD_FOLDER'], filename)
    
    df = pd.DataFrame(contacts)
    df = df.rename(columns={
        'name': 'שם',
        'phone': 'טלפון',
        'email': 'אימייל',
        'address': 'כתובת',
        'source': 'מקור'
    })
    
    writer = pd.ExcelWriter(output_path, engine='xlsxwriter')
    df.to_excel(writer, index=False, sheet_name='אנשי קשר')
    
    # עיצוב הקובץ
    workbook = writer.book
    worksheet = writer.sheets['אנשי קשר']
    
    header_format = workbook.add_format({
        'bold': True,
        'align': 'center',
        'valign': 'vcenter',
        'bg_color': '#4F81BD',
        'font_color': 'white',
        'border': 1
    })
    
    cell_format = workbook.add_format({
        'align': 'right',
        'valign': 'vcenter',
        'border': 1
    })
    
    # הגדרת רוחב עמודות ועיצוב
    for col_num, value in enumerate(df.columns.values):
        worksheet.write(0, col_num, value, header_format)
        max_length = max(df[value].astype(str).apply(len).max(), len(value))
        worksheet.set_column(col_num, col_num, max_length + 2, cell_format)
    
    writer.close()
    return output_path

@app.route('/download/<filename>')
def download_file(filename):
    """הורדת קובץ התוצאה"""
    try:
        return send_file(
            os.path.join(app.config['DOWNLOAD_FOLDER'], filename),
            as_attachment=True,
            download_name=f'אנשי_קשר_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
        )
    except Exception as e:
        logger.error(f"שגיאה בהורדת הקובץ: {str(e)}")
        return jsonify({'success': False, 'error': 'שגיאה בהורדת הקובץ'})

def is_port_available(port):
    """בודק אם הפורט פנוי"""
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        try:
            s.bind(('0.0.0.0', port))
            return True
        except OSError:
            return False

if __name__ == '__main__':
    try:
        port = 5001
        while port < 5100:
            if is_port_available(port):
                break
            port += 1
        
        if port >= 5100:
            logger.error("לא נמצא פורט פנוי בטווח 5001-5099")
            exit(1)
            
        logger.info(f"מפעיל את השרת בפורט {port}")
        print(f"גש לאפליקציה בכתובת: http://localhost:{port}")
        
        app.run(host='0.0.0.0', port=port, debug=True)
    except Exception as e:
        logger.error(f"שגיאה בהפעלת השרת: {str(e)}") 