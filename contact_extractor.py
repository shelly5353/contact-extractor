import os
import re
import logging
import json
import pandas as pd
import PyPDF2
import pdfplumber
import openpyxl
import docx2txt
from typing import Dict, List, Set, Optional, Tuple
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl.styles import PatternFill, Font, Alignment
from openpyxl.utils import get_column_letter
from PyPDF2 import PdfReader
from pdf2image import convert_from_path
import pytesseract
import subprocess
import tempfile
import shutil
import olefile
from docx import Document
import ssl
import urllib.request
from tika import parser
import certifi
import urllib3

# Configure logging
logging.basicConfig(
    format='%(asctime)s - %(levelname)s - %(message)s',
    level=logging.INFO
)

class ContactExtractor:
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.phone_pattern = re.compile(r'(?:\+972|05|972|0)[0-9\-\s]{8,}')
        self.email_pattern = re.compile(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}')
        self.address_pattern = r'(?:רחוב|רח\'|שד\'|שדרות|דרך|כביש|סמטת|סמ\')\s+[\u0590-\u05FF\s\w.,\-\'"]+'
        
        # Hebrew name pattern with titles and variations
        self.title_pattern = r'(?:ד"ר|דר\'|ד״ר|עו"ד|עו״ד|פרופ\'|פרופסור|גב\'|גברת|מר|רו"ח|רו״ח|דוקטור|אינג\'|אדר\'|אדריכל|מהנדס)'
        self.name_pattern = re.compile(r'[א-ת\s]{2,}')
        
        # Pattern for lines that likely contain a name
        self.name_line_pattern = f'(?:{self.title_pattern}\\s+)?[\\u0590-\\u05FF]+(?:\\s+[\\u0590-\\u05FF]+){{1,3}}(?:\\s*[-,]\\s*[^\\n]*)?'
        
        # Pattern for role/position after name
        self.role_suffix_pattern = r'(?:\s*[-,]\s*(?:מנהל\S*|יועץ\S*|אחראי\S*|רכז\S*|עובד\S*|ראש\s+\S+|סגן\S*|מפקח\S*|מהנדס\S*|אדריכל\S*|מתכנן\S*).*)?'
        
        # Israeli address pattern with variations
        self.street_prefixes = r'(?:רח\'|רחוב|שד\'|שדרות|דרך|סמטת|סמ\'|שכונת|שכ\')'
        self.city_names = r'(?:תל[- ]אביב|רמת[- ]גן|חיפה|ירושלים|באר[- ]שבע|רחובות|פתח[- ]תקווה|רעננה|הרצליה|גבעתיים|חולון|בת[- ]ים|נתניה|אשדוד|אשקלון|רמלה|לוד|כפר[- ]סבא|רמת[- ]השרון|ראש[- ]העין|פ"ת|פ״ת)'
        
        # Words that indicate a role/title rather than a name
        self.role_words = [
            'מנהל', 'מנהלת', 'יועץ', 'יועצת', 'עובד', 'עובדת', 'אחראי', 'אחראית',
            'מזכיר', 'מזכירה', 'ראש', 'סגן', 'סגנית', 'מפקח', 'מפקחת', 'רכז', 'רכזת',
            'משפטי', 'משפטית', 'כספים', 'פרויקט', 'שירות', 'לקוחות', 'תפעול', 'מחקר',
            'פיתוח', 'משאבי', 'אנוש', 'סניף', 'מחלקה', 'צוות', 'גיוס', 'סוכנים',
            'מכירות', 'שיווק', 'תמיכה', 'הדרכה', 'אזור', 'מרכז', 'צפון', 'דרום',
            'מערב', 'מזרח', 'סוכן', 'סוכנת', 'נציג', 'נציגה', 'מוקד', 'שלוחה',
            'בטיפול', 'במעקב', 'הועבר', 'לא', 'כרגע', 'רלוונטי', 'נרשם', 'נרשמה',
            'מעוניין', 'ניסיתי', 'רחוק', 'קורס', 'לימודי', 'בוקר', 'השקעות'
        ]
        
        # Contact label patterns
        self.contact_labels = [
            'טלפון', 'נייד', 'טל', 'פקס', 'דוא"ל', 'אימייל', 'מייל', 'כתובת',
            'שם', 'איש קשר', 'פרטי התקשרות', 'פרטים', 'פרטי קשר', 'תפקיד',
            'משרד', 'סניף', 'מחלקה', 'יחידה', 'אגף', 'מטה', 'הנהלה'
        ]
        
        # Common words that should not be treated as names
        self.non_name_words = [
            'גיוס', 'סוכנים', 'אזור', 'מרכז', 'צפון', 'דרום', 'מערב', 'מזרח',
            'ראשי', 'משני', 'חדש', 'ישן', 'כללי', 'מיוחד', 'נוסף', 'אחר',
            'ראשון', 'שני', 'שלישי', 'רביעי', 'חמישי', 'שישי', 'שביעי',
            'מחוז', 'מרחב', 'שכונה', 'רובע', 'קומה', 'בניין', 'מתחם',
            'ראש העין', 'פתח תקווה', 'תל אביב', 'רמת גן', 'גבעתיים',
            'הרצליה', 'רעננה', 'כפר סבא', 'נתניה', 'חולון', 'בת ים',
            'ירושלים', 'חיפה', 'באר שבע', 'אשדוד', 'אשקלון',
            'בטיפול', 'במעקב', 'הועבר', 'לא', 'כרגע', 'רלוונטי',
            'נרשם', 'נרשמה', 'מעוניין', 'ניסיתי', 'רחוק', 'קורס',
            'לימודי', 'בוקר', 'השקעות', 'פ"ת', 'פ״ת'
        ]
        
        # Israeli phone patterns with variations
        self.phone_prefixes = r'(?:טלפון|נייד|טל\'?|פקס|סלולרי|נייח|מספר|חירום|מוקד|שלוחה)'
        self.phone_patterns = [
            f'{self.phone_prefixes}?\\s*:?\\s*(?:\\+972-?|972-?|0)(?:[23489]|5[0-9]|7[0-9]|81)-?[2-9]\\d{{2}}-?\\d{{4}}',
            r'(?:\+972-?|972-?|0)(?:[23489]|5[0-9]|7[0-9]|81)-?[2-9]\d{2}-?\d{4}',  # Without prefix
            r'(?:\+972-?|972-?|0)(?:[23489]|5[0-9]|7[0-9]|81)(?:\s*-?\s*\d{3}){2,3}',  # Spaced groups
            r'\d{2,3}[-\s]?\d{3}[-\s]?\d{4}'  # Simple format
        ]

        # Initialize contact storage
        self.contacts = []
        self.processed_files = set()

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Replace various unicode quotation marks and apostrophes
        text = re.sub(r'[""״\'׳]', '"', text)
        # Replace multiple spaces and newlines
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    def _extract_text_from_table(self, table: Table) -> List[str]:
        """Extract text from table cells."""
        text_lines = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                # Check if this is a contact row (should have 3 cells: name, city, phone)
                if len(cells) == 3:
                    # Join with a special separator that we can split on later
                    text_lines.append(" ||| ".join(cells))
                else:
                    # For other rows, just join with a space
                    text_lines.append(" ".join(cells))
        return text_lines

    def _is_likely_name_line(self, line: str) -> bool:
        """Check if a line is likely to contain a name."""
        # Skip empty lines or lines that are too short
        if not line or len(line) < 3:
            return False
            
        # Skip lines that are too long (likely not a name)
        if len(line.split()) > 8:  # Increased to allow for titles and roles
            return False
            
        # Skip lines that look like addresses or contain numbers
        if (re.search(self.street_prefixes, line) or 
            re.search(r'\d{3,}', line) or
            re.search(r'קומה|דירה|כניסה|בניין|מתחם', line)):
            return False
            
        # Skip lines that contain contact information
        if any(pattern in line.lower() for pattern in [
            'טלפון:', 'נייד:', 'טל:', 'פקס:', 'דוא"ל:', 'אימייל:', 'מייל:', 'כתובת:',
            'ת.ד.', 'מיקוד', 'ת.ז.', '@', 'שלוחה', 'מוקד:', 'אזור', 'ראש העין'
        ]):
            return False
            
        # Skip lines that are only role words or non-name words
        words = line.split()
        non_name_words = [w for w in words if w not in self.role_words and 
                         w not in self.contact_labels and 
                         w not in self.non_name_words]
        if not non_name_words:
            return False
            
        # Check if line starts with a title or contains name indicators
        if re.match(self.title_pattern, line) or any(indicator in line for indicator in ['שם:', 'איש קשר:', 'נציג:']):
            return True
            
        # Check if line contains Hebrew name pattern (at least two Hebrew words)
        if re.search(r'[\u0590-\u05FF]+(?:\s+[\u0590-\u05FF]+){1,3}', line):
            # Additional check for common name formats
            if re.search(r'(?:^|\s)(?:מר|גב\'|ד"ר|עו"ד|רו"ח)\s+[\u0590-\u05FF]+', line):
                return True
            # Check for name followed by role
            if re.search(r'[\u0590-\u05FF]+(?:\s+[\u0590-\u05FF]+){1,2}\s*[-,]\s*[^,\n]+$', line):
                # Make sure the first part doesn't contain non-name words
                first_part = line.split('[-,]')[0]
                if not any(word in first_part for word in self.non_name_words):
                    return True
            # Check for standalone name
            if (len(line.split()) <= 4 and 
                all(re.search(r'[\u0590-\u05FF]', word) for word in line.split()) and
                not any(word in line for word in self.non_name_words + ['אזור', 'ראש העין'])):
                return True
            
        return False

    def _extract_name_from_line(self, line: str) -> Optional[str]:
        """Extract a name from a line of text."""
        # Remove status indicators and comments
        line = re.sub(r'^(?:בטיפול|במעקב|הועבר ל|לא כרגע|לא רלוונטי|נרשמה?|מעוניין|ניסיתי|רחוק)\s*-?\s*', '', line)
        line = re.sub(r'\s*-\s*.*$', '', line)  # Remove everything after a dash
        
        # Remove common prefixes and labels
        line = re.sub(r'^(?:שם|איש קשר|נציג|עובד|פרטי|תפקיד|מייל|טלפון)\s*:?\s*', '', line)
        line = re.sub(r'^(?:לכבוד|לידי|עבור)\s+', '', line)
        
        # Clean up the line
        line = line.strip()
        if not line:
            return None
            
        # Try to match the name pattern
        match = re.match(self.name_line_pattern, line)
        if not match:
            return None
            
        # Get the matched text
        name = match.group(0)
        
        # Handle cases with commas or dashes
        parts = re.split(r'[-,]', name)
        name = parts[0].strip()
        
        # If we have additional parts that look like part of the name, include them
        if len(parts) > 1:
            second_part = parts[1].strip()
            # Check if the second part looks like a name continuation
            if (len(second_part.split()) <= 2 and 
                re.search(r'[\u0590-\u05FF]', second_part) and
                not any(word in second_part for word in self.role_words)):
                name = f"{name} {second_part}"
        
        # Skip if it's just a title or common word
        if name in ['מר', 'גב\'', 'ד"ר', 'פרופ\'', 'עו"ד', 'רו"ח', 'אדון', 'גברת']:
            return None
            
        # Skip if it contains only role words or common words
        words = name.split()
        non_role_words = [w for w in words if w not in self.role_words and w not in self.contact_labels]
        if not non_role_words:
            return None
            
        # Verify the name contains Hebrew characters and meets length requirements
        if not re.search(r'[\u0590-\u05FF]', name) or len(name) < 3:
            return None
            
        # Verify maximum number of words and minimum word length
        name_words = name.split()
        if (len(name_words) > 4 or 
            any(len(word) < 2 for word in name_words if re.search(r'[\u0590-\u05FF]', word))):
            return None
            
        # Skip if it's a heading or common phrase
        if name.lower() in [
            'רשימת אנשי קשר', 'אנשי קשר', 'פרטי קשר', 'פרטי התקשרות',
            'מחלקת שירות', 'צוות פיתוח', 'הנהלת חשבונות'
        ]:
            return None
            
        return name

    def _process_name_and_role(self, name: str, role: str) -> str:
        """Process a name and role combination."""
        # Clean up the role
        role = role.strip()
        role = re.sub(r'^[,-]\s*', '', role)  # Remove leading comma or dash
        role = re.sub(r'^\s*(?:תפקיד|פקיד)\s*:?\s*', '', role)  # Remove role label
        
        # Check if the role is valid and not contact information
        if role and not any(word in role.lower() for word in [
            'טלפון', 'נייד', 'פקס', 'דוא"ל', 'מייל', 'כתובת', '@', 
            'ת.ד.', 'מיקוד', 'ת.ז.'
        ]):
            # Clean up role format
            role = re.sub(r'\s+', ' ', role)  # Normalize spaces
            role = re.sub(r'[.]{2,}', '', role)  # Remove multiple dots
            return f"{name} - {role}"
        return name

    def _extract_name_and_role(self, text: str) -> Optional[Tuple[str, Optional[str]]]:
        """Extract name and role from text."""
        # Try to find name and role in format: "name - role"
        match = re.match(r'^(.+?)\s*[-,]\s*(.+?)(?=(?:,|\s*(?:טלפון|נייד|טל|פקס|דוא"ל|אימייל|מייל|כתובת)\s*:|$))', text)
        if match:
            name = match.group(1).strip()
            role = match.group(2).strip()
            if self._is_likely_name_line(name):
                name = self._extract_name_from_line(name)
                if name:
                    return name, role
        
        # Try to find just a name
        if self._is_likely_name_line(text):
            name = self._extract_name_from_line(text)
            if name:
                return name, None
        
        return None

    def extract_phone_numbers(self, text: str) -> List[str]:
        """Extract phone numbers from text using regex patterns that match the original format."""
        phone_numbers = set()
        text = self._clean_text(text)
        
        # Define patterns that match the exact format in the file
        patterns = [
            r'0[23489]-\d{7}',  # Format: 03-9130953
            r'05[0-9]-\d{7}'    # Format: 054-7599959
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                phone_numbers.add(match.group())
                
        return sorted(list(phone_numbers))

    def _is_valid_email(self, email: str) -> bool:
        """Validate email address."""
        # Basic validation
        if not re.match(r'^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}$', email):
            return False
            
        # Check for common invalid patterns
        invalid_patterns = [
            r'example\.com$',
            r'test\.com$',
            r'domain\.com$',
            r'@.*@',
            r'\.{2,}',
            r'^[0-9]+@'
        ]
        
        for pattern in invalid_patterns:
            if re.search(pattern, email):
                return False
                
        # Check for valid domains
        domain = email.split('@')[1]
        valid_domains = [
            'gmail.com', 'yahoo.com', 'hotmail.com', 'walla.co.il',
            'gmail.co.il', 'bezeqint.net', 'zahav.net.il', 'netvision.net.il',
            'outlook.com', 'aol.com', 'mail.com', 'office.com'
        ]
        
        if '.' not in domain:
            return False
            
        return True

    def _extract_emails(self, text: str) -> List[str]:
        """Extract email addresses from text."""
        emails = set()
        
        # Basic email pattern
        email_pattern = r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}'
        
        # Find all matches
        matches = re.finditer(email_pattern, text)
        for match in matches:
            email = match.group().lower()
            # Additional validation
            if self._is_valid_email(email):
                emails.add(email)
                
        # Look for emails with Hebrew text around them
        hebrew_context_pattern = r'[\u0590-\u05FF\s]+({})[\u0590-\u05FF\s]+'.format(email_pattern)
        matches = re.finditer(hebrew_context_pattern, text)
        for match in matches:
            email = match.group(1).lower()
            if self._is_valid_email(email):
                emails.add(email)
        
        return sorted(list(emails))

    def extract_names(self, text: str) -> List[str]:
        """Extract potential names from text."""
        names = set()
        
        # Split text into lines and process each line
        for line in text.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            # Check if this is a table row with our special separator
            parts = line.split(" ||| ")
            if len(parts) == 3:  # name, city, phone
                name = parts[0].strip()
                if name and not any(word in name.lower() for word in self.non_name_words):
                    names.add(name)
            else:
                # Process regular lines
                result = self._extract_name_and_role(line)
                if result:
                    name, role = result
                    names.add(name)
        
        return sorted(list(names))

    def extract_addresses(self, text: str) -> List[str]:
        """Extract potential addresses from text."""
        text = self._clean_text(text)
        addresses = set()
        
        # Look for addresses in each line
        for line in text.split('\n'):
            line = line.strip()
            if not line:
                continue
            
            # Try to find addresses in the line
            matches = re.finditer(self.address_pattern, line)
            for match in matches:
                # Get the full line from the match to the end
                addr_text = line[match.start():].strip()
                
                # Try to find the city
                city_match = re.search(f',\\s*({self.city_names})', addr_text)
                if city_match:
                    # Take everything up to the end of the city name
                    addr = addr_text[:city_match.end()].strip()
                else:
                    # Take the first part until a clear separator
                    addr = addr_text.split('|')[0].split('נייד')[0].split('טל')[0].strip()
                    # Split on common separators and take the first part
                    separators = ['מסביב', 'ליד', 'ניתן', 'שעות', 'פתוח', 'חניה', 'כניסה חופשית', 'עלות', '₪', 'ש"ח']
                    for sep in separators:
                        if sep in addr:
                            addr = addr.split(sep)[0].strip()
                    addr = re.split(r'\s{2,}', addr)[0].strip()
                
                # Clean up the address
                addr = re.sub(r'^(?:כתובת|משרד|סניף|למשלוח דואר)\s*:?\s*', '', addr)
                addr = re.sub(r'^\s*:\s*', '', addr)
                
                # Remove any trailing punctuation or common words
                addr = re.sub(r'[.,:;]+$', '', addr)
                addr = re.sub(r'\s+(?:בלבד|בחינם|חינם|חופשי|נוסף)$', '', addr)
                
                # Verify the address looks valid
                if (addr and 
                    re.search(r'\d+', addr) and  # Must contain a number
                    not re.search(r'₪|ש"ח|\$', addr) and  # Should not contain currency
                    len(addr.split()) >= 3):  # Must have at least 3 words
                    addresses.add(addr)
        
        return sorted(list(addresses))

    def _normalize_name(self, name: str) -> str:
        """Normalize a name by sorting its parts to handle cases like 'טל עוזר' and 'עוזר טל'."""
        # Split the name into parts and sort them
        parts = name.split()
        return ' '.join(sorted(parts))

    def extract_from_docx(self, file_path: str) -> Dict[str, List[str]]:
        """Extract all contact information from a Word document."""
        try:
            doc = Document(file_path)
            contacts_dict = {}  # Dictionary to store unique contacts by phone number
            has_valid_contacts = False
            
            # Process all tables
            for table in doc.tables:
                # First, try to identify columns that might contain contact information
                header_row = table.rows[0] if table.rows else None
                if not header_row:
                    continue
                
                # Get all header texts
                headers = [cell.text.strip().lower() for cell in header_row.cells]
                
                # Try to identify relevant columns
                name_col = -1
                phone_col = -1
                city_col = -1
                
                for idx, header in enumerate(headers):
                    # Look for name column
                    if any(word in header for word in ['שם', 'איש קשר', 'נציג']):
                        name_col = idx
                    # Look for phone column
                    elif any(word in header for word in ['טלפון', 'נייד', 'טל', 'פלאפון']):
                        phone_col = idx
                    # Look for city/address column
                    elif any(word in header for word in ['עיר', 'ישוב', 'כתובת', 'מיקום']):
                        city_col = idx
                
                # If we couldn't identify columns by headers, try to identify by content
                if name_col == -1 or phone_col == -1:
                    # Skip header row and check content
                    for row in table.rows[1:]:
                        for idx, cell in enumerate(row.cells):
                            content = cell.text.strip()
                            # Look for phone number pattern
                            if re.match(r'0[23489]-\d{7}|05[0-9]-\d{7}', content):
                                phone_col = idx
                            # Look for name pattern (2-3 Hebrew words)
                            elif (re.match(r'[\u0590-\u05FF]+(?:\s+[\u0590-\u05FF]+){1,2}$', content) and
                                  not any(word in content.lower() for word in self.non_name_words)):
                                name_col = idx
                
                # Process rows if we found at least name and phone columns
                if name_col != -1 and phone_col != -1:
                    for row_idx, row in enumerate(table.rows):
                        if row_idx == 0:  # Skip header row
                            continue
                            
                        cells = [cell.text.strip() for cell in row.cells]
                        if len(cells) <= max(name_col, phone_col):
                            continue
                            
                        name = cells[name_col]
                        phone = cells[phone_col].replace('\t', '')
                        city = cells[city_col] if city_col != -1 and city_col < len(cells) else ''
                        
                        # Skip empty cells
                        if not name or not phone:
                            continue
                            
                        # Additional validation for name
                        if (len(name) < 3 or  # Name too short
                            any(word in name.lower() for word in self.non_name_words) or  # Contains non-name words
                            not re.search(r'[\u0590-\u05FF]', name)):  # No Hebrew characters
                            continue
                            
                        # Validate phone number format
                        if re.match(r'0[23489]-\d{7}|05[0-9]-\d{7}', phone):
                            has_valid_contacts = True
                            # Use the normalized name for comparison
                            norm_name = self._normalize_name(name)
                            
                            # If we already have this phone number, keep the longer name
                            if phone in contacts_dict:
                                existing_name = contacts_dict[phone][0]
                                if len(name) > len(existing_name):
                                    contacts_dict[phone] = (name, city)
                            else:
                                contacts_dict[phone] = (name, city)
            
            if not has_valid_contacts:
                print("לא נמצאו אנשי קשר תקינים במסמך זה.")
                return {
                    "names": [],
                    "phones": [],
                    "emails": [],
                    "addresses": []
                }
            
            # Convert the dictionary to sorted lists
            sorted_contacts = sorted(contacts_dict.items(), key=lambda x: x[1][0])  # Sort by name
            names = []
            phones = []
            cities = []
            
            # Print the results and build the lists
            print("\nאנשי קשר שנמצאו במסמך:")
            print("=" * 60)
            print(f"{'שם':<30} {'טלפון':<15} {'עיר'}")
            print("-" * 60)
            
            for phone, (name, city) in sorted_contacts:
                print(f"{name:<30} {phone:<15} {city}")
                names.append(name)
                phones.append(phone)
                cities.append(city)
            
            print("=" * 60)
            print(f"סה\"כ נמצאו {len(names)} אנשי קשר")
            print()
            
            return {
                "names": names,
                "phones": phones,
                "emails": [],
                "addresses": cities,
            }
            
        except Exception as e:
            print(f"שגיאה בעיבוד המסמך {os.path.basename(file_path)}: {str(e)}")
            return {
                "names": [],
                "phones": [],
                "emails": [],
                "addresses": []
            }

    def _extract_text_from_old_doc(self, file_path):
        try:
            # ניסיון להמיר את הקובץ ל-PDF באמצעות LibreOffice ואז לתמונה באמצעות ImageMagick
            try:
                print(f"מנסה להמיר את הקובץ {os.path.basename(file_path)} באמצעות LibreOffice ו-OCR...")
                
                # המרה ל-PDF
                pdf_path = file_path + '.pdf'
                cmd = ['/Applications/LibreOffice.app/Contents/MacOS/soffice', '--headless', '--convert-to', 'pdf', 
                       '--outdir', os.path.dirname(file_path), file_path]
                subprocess.run(cmd, capture_output=True, check=True)
                
                # בדיקה אם ה-PDF נוצר
                pdf_expected_path = os.path.join(os.path.dirname(file_path), os.path.basename(file_path) + '.pdf')
                if os.path.exists(pdf_expected_path):
                    # המרת ה-PDF לתמונה
                    png_path = file_path + '.png'
                    convert_path = '/opt/homebrew/bin/convert'
                    cmd = [convert_path, '-density', '300', pdf_expected_path, png_path]
                    subprocess.run(cmd, capture_output=True, check=True)
                    
                    # בדיקה אם התמונה נוצרה
                    if os.path.exists(png_path):
                        try:
                            # בדיקת השפות הזמינות
                            langs_output = subprocess.run(['tesseract', '--list-langs'], 
                                                        capture_output=True, text=True)
                            available_langs = langs_output.stdout.strip().split('\n')[1:]
                            
                            # הגדרת השפות לשימוש
                            lang_param = 'eng'  # ברירת מחדל - אנגלית
                            
                            # אם עברית זמינה, נוסיף אותה
                            if 'heb' in available_langs:
                                lang_param = 'heb+eng'
                            
                            # חילוץ טקסט מהתמונה
                            text = subprocess.run(['tesseract', png_path, 'stdout', '-l', lang_param], 
                                                capture_output=True, text=True, encoding='utf-8').stdout
                            
                            # ניקוי קבצים זמניים
                            try:
                                os.unlink(png_path)
                                os.unlink(pdf_expected_path)
                            except Exception as e:
                                print(f"שגיאה בניקוי קבצים זמניים: {str(e)}")
                            
                            if text and any(c.isalnum() for c in text):
                                print(f"OCR הצליח לחלץ טקסט מהקובץ {os.path.basename(file_path)}")
                                return text
                        except Exception as e:
                            print(f"OCR text extraction failed: {str(e)}")
                            # ניקוי קבצים זמניים
                            try:
                                if os.path.exists(png_path):
                                    os.unlink(png_path)
                                if os.path.exists(pdf_expected_path):
                                    os.unlink(pdf_expected_path)
                            except Exception as e:
                                print(f"שגיאה בניקוי קבצים זמניים: {str(e)}")
            except Exception as e:
                print(f"LibreOffice/ImageMagick conversion failed: {str(e)}")
            
            # אם ההמרה נכשלה, ננסה את שאר השיטות
            # כלי לתיקון קבצי DOC פגומים
            try:
                # יצירת עותק של הקובץ המקורי למקרה שהתיקון יפגע בו
                backup_path = file_path + '.backup'
                shutil.copy2(file_path, backup_path)
                
                # ניסיון לתקן את הקובץ (הוצאת מידע מהקובץ המקורי וייצור קובץ חדש)
                recovered_path = file_path + '.recovered.docx'
                
                # ניסיון לשחזר באמצעות דוק אקס פי
                try:
                    with open(backup_path, 'rb') as infile, open(recovered_path, 'wb') as outfile:
                        # קריאת ה-bytes של הקובץ
                        content = infile.read()
                        
                        # חיפוש סמנים של פורמט דוקס
                        docx_marker = b'PK\x03\x04'
                        docx_pos = content.find(docx_marker)
                        
                        if docx_pos >= 0:
                            # יש סממנים של קובץ DOCX בתוך קובץ ה-DOC - נחלץ אותו
                            outfile.write(content[docx_pos:])
                            print(f"נמצא קובץ DOCX מוטמע בתוך קובץ DOC, נשמר ב-{recovered_path}")
                            
                            # נסה לקרוא את הקובץ המשוחזר
                            try:
                                doc = Document(recovered_path)
                                text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                                if text and any(c.isalnum() for c in text):
                                    return text
                            except Exception as e:
                                print(f"Failed to read recovered docx: {str(e)}")
                except Exception as e:
                    print(f"DOC repair attempt failed: {str(e)}")
                finally:
                    # מחיקת קבצים זמניים
                    if os.path.exists(recovered_path):
                        try:
                            os.unlink(recovered_path)
                        except:
                            pass
                    if os.path.exists(backup_path):
                        try:
                            os.unlink(backup_path)
                        except:
                            pass
            except Exception as e:
                print(f"Document repair attempt failed: {str(e)}")
            
            # ניסיון עם PyPDF2 (למקרה שהקובץ הוא למעשה PDF בתחפושת)
            try:
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    text = ''
                    for page in pdf_reader.pages:
                        text += page.extract_text() + '\n'
                    if text.strip() and any(c.isalnum() for c in text):
                        return text
            except Exception as e:
                print(f"PyPDF2 attempt failed: {str(e)}")
                
            # ניסיון עם tika (Apache Tika)
            try:
                from tika import parser
                parsed = parser.from_file(file_path)
                text = parsed["content"]
                if text and any(c.isalnum() for c in text):
                    # חיפוש חלקים בעברית
                    hebrew_parts = re.findall(r'[A-Za-z0-9\s]*[\u0590-\u05FF]{2,}[^\n]*', text)
                    if hebrew_parts:
                        return '\n'.join(hebrew_parts)
                    return text
            except Exception as e:
                print(f"tika attempt failed: {str(e)}")
            
            # נסיון עם docx2python
            try:
                result = docx2python.docx2python(file_path)
                text = '\n'.join(result.text)
                
                if text and any(c.isalnum() for c in text):
                    return text
            except Exception as e:
                print(f"docx2python attempt failed: {str(e)}")
            
            # נתיב מלא לקובץ
            absolute_path = os.path.abspath(file_path)
            
            # ניסיון להשתמש ב-pywin32 אם זמין (בד"כ זמין רק בWindows או דרך Wine)
            try:
                import win32com.client
                import pythoncom
                
                # אתחול COM
                pythoncom.CoInitialize()
                
                # יצירת אובייקט Word
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                
                # פתיחת המסמך
                doc = word.Documents.Open(absolute_path)
                
                # חילוץ הטקסט
                text = doc.Content.Text
                
                # סגירת המסמך ו-Word
                doc.Close()
                word.Quit()
                
                if text and any(c.isalnum() for c in text):
                    return text
            except Exception as e:
                print(f"pywin32 attempt failed: {str(e)}")
            
            # ניסיון להשתמש בTexttospeech (כלי של macOS) - לא עובד עם .doc
            try:
                with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as temp_file:
                    temp_file_path = temp_file.name
                
                cmd = ['textutil', '-convert', 'txt', '-output', temp_file_path, file_path]
                subprocess.run(cmd, check=True, capture_output=True)
                
                with open(temp_file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                
                os.unlink(temp_file_path)
                
                if text and any(c.isalnum() for c in text):
                    return text
            except Exception as e:
                print(f"textutil attempt failed: {str(e)}")
            
            # ניסיון להשתמש ב-hexdump וחיפוש טקסטים בתוכו
            try:
                cmd = ['hexdump', '-C', file_path]
                result = subprocess.run(cmd, capture_output=True, text=True, check=True)
                
                # חיפוש מחרוזות עברית בפלט ההקסדציימלי
                hebrew_patterns = []
                output = result.stdout
                
                # חיפוש מילים בעברית בפלט הבינארי (חלקי)
                hebrew_words = re.findall(r'[\u0590-\u05FF]{2,}', output)
                if hebrew_words:
                    return ' '.join(hebrew_words)
            except Exception as e:
                print(f"hexdump attempt failed: {str(e)}")
            
            # ניסיון להשתמש ב-strings עם אפשרויות שונות לזיהוי עברית
            encodings = ['UTF-8', 'windows-1255', 'hebrew', 'cp1255', 'iso-8859-8']
            try:
                # הרצת strings עם אפשרויות שונות
                for encoding in encodings:
                    cmd = ['strings', file_path]
                    result = subprocess.run(cmd, capture_output=True, text=True, check=True)
                    
                    try:
                        # ניסיון להמיר את הפלט לקידוד המתאים
                        output = result.stdout.encode().decode(encoding, errors='ignore')
                        
                        # אם יש טקסט בעברית, החזר אותו
                        if re.search(r'[\u0590-\u05FF]{2,}', output):
                            return output
                    except Exception:
                        continue
            except Exception as e:
                print(f"strings attempt failed: {str(e)}")
            
            # ניסיון עם חילוץ מ-OLE עם זיהוי אוטומטי של קידוד
            try:
                ole = olefile.OleFileIO(file_path)
                
                # ניסיון למצוא זרמים רלוונטיים
                relevant_streams = []
                for stream in ole.listdir():
                    stream_path = '/'.join(stream)
                    if any(x in stream_path.lower() for x in ['word', 'document', 'text', 'content', '1table']):
                        relevant_streams.append(stream)
                
                # אם לא מצאנו זרמים רלוונטיים, בדוק את כל הזרמים
                if not relevant_streams:
                    relevant_streams = ole.listdir()
                
                # עבור כל זרם, נסה לחלץ תוכן
                for stream in relevant_streams:
                    try:
                        with ole.openstream(stream) as stream_obj:
                            content = stream_obj.read()
                            
                            # נסה עם כל קידוד אפשרי
                            for encoding in encodings:
                                try:
                                    text = content.decode(encoding, errors='ignore')
                                    # בדוק אם יש תוכן עברי משמעותי
                                    if re.search(r'[\u0590-\u05FF]{2,}', text):
                                        # מצא רק חלקים בעברית עם תווים סביבם
                                        hebrew_parts = re.findall(r'[A-Za-z0-9\s]*[\u0590-\u05FF]{2,}[^\n]*', text)
                                        if hebrew_parts:
                                            return '\n'.join(hebrew_parts)
                                except Exception:
                                    continue
                    except Exception:
                        continue
            except Exception as e:
                print(f"Enhanced olefile attempt failed: {str(e)}")
            
            # נסה להשתמש ב-abiword אם מותקן
            try:
                with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp:
                    tmp_path = tmp.name
                
                cmd = ['abiword', '--to=txt', '--to-name=' + tmp_path, file_path]
                subprocess.run(cmd, capture_output=True, check=True)
                
                with open(tmp_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                
                os.unlink(tmp_path)
                
                if text and any(c.isalnum() for c in text):
                    return text
            except Exception as e:
                print(f"abiword attempt failed: {str(e)}")
            
            # נסה להשתמש ב-unrtf
            try:
                cmd = ['unrtf', '--text', file_path]
                result = subprocess.run(cmd, capture_output=True, text=True, check=True)
                
                if result.stdout.strip():
                    # נסה להמיר לקידודים שונים
                    for encoding in encodings:
                        try:
                            text = result.stdout.encode().decode(encoding, errors='ignore')
                            if re.search(r'[\u0590-\u05FF]{2,}', text):
                                return text
                        except Exception:
                            continue
            except Exception as e:
                print(f"unrtf attempt failed: {str(e)}")
            
            # Try oletools first
            try:
                from oletools.olevba import VBA_Parser
                vba_parser = VBA_Parser(file_path)
                
                # Try to read all streams
                ole = vba_parser.ole_file
                if ole:
                    # Try to read WordDocument stream
                    if ole.exists('WordDocument'):
                        with ole.openstream('WordDocument') as stream:
                            content = stream.read()
                            # Try different encodings
                            for encoding in ['utf-16le', 'utf-8', 'cp1255', 'windows-1255', 'iso-8859-8', 'hebrew']:
                                try:
                                    text = content.decode(encoding, errors='ignore')
                                    if text and any(c.isalnum() for c in text):
                                        return text
                                except Exception:
                                    continue
                    
                    # Try to read other common streams
                    for stream_name in ['1Table', 'Table', 'Data', 'ObjectPool']:
                        if ole.exists(stream_name):
                            with ole.openstream(stream_name) as stream:
                                content = stream.read()
                                # Try different encodings
                                for encoding in ['utf-16le', 'utf-8', 'cp1255', 'windows-1255', 'iso-8859-8', 'hebrew']:
                                    try:
                                        text = content.decode(encoding, errors='ignore')
                                        if text and any(c.isalnum() for c in text):
                                            return text
                                    except Exception:
                                        continue
            except Exception as e:
                print(f"oletools attempt failed: {str(e)}")

            # Try strings
            try:
                cmd = ['strings', '-e', 'l', file_path]  # -e l for little-endian 16-bit characters
                result = subprocess.run(cmd, capture_output=True, text=True, check=True)
                if result.stdout.strip():
                    return result.stdout
            except Exception as e:
                print(f"strings attempt failed: {str(e)}")

            # Try soffice (LibreOffice)
            try:
                # Convert .doc to .txt using soffice
                txt_path = os.path.splitext(file_path)[0] + '.txt'
                cmd = ['/Applications/LibreOffice.app/Contents/MacOS/soffice', '--headless', '--convert-to', 'txt:Text', file_path]
                result = subprocess.run(cmd, capture_output=True, text=True, check=True)
                
                # Read the converted text file
                if os.path.exists(txt_path):
                    with open(txt_path, 'r', encoding='utf-8') as f:
                        text = f.read()
                    os.unlink(txt_path)  # Clean up
                    if text and any(c.isalnum() for c in text):
                        return text
            except Exception as e:
                print(f"soffice attempt failed: {str(e)}")

            # Try unoconv
            try:
                # Convert .doc to .txt using unoconv
                txt_path = file_path + '.txt'
                cmd = ['unoconv', '--format=text', '--output=' + txt_path, file_path]
                result = subprocess.run(cmd, capture_output=True, text=True, check=True)
                
                # Read the converted text file
                if os.path.exists(txt_path):
                    with open(txt_path, 'r', encoding='utf-8') as f:
                        text = f.read()
                    os.unlink(txt_path)  # Clean up
                    if text and any(c.isalnum() for c in text):
                        return text
            except Exception as e:
                print(f"unoconv attempt failed: {str(e)}")

            # Try pandoc
            try:
                cmd = ['pandoc', '-f', 'doc', '-t', 'plain', file_path]
                result = subprocess.run(cmd, capture_output=True, text=True, check=True)
                if result.stdout.strip():
                    return result.stdout
            except Exception as e:
                print(f"pandoc attempt failed: {str(e)}")

            # Try wv with different encodings and options
            encodings = ['UTF-8', 'windows-1255', 'hebrew', 'cp1255', 'iso-8859-8']
            wv_options = [
                ['wvText', file_path, '-'],  # Default
                ['wvText', file_path, '-', '--charset=utf-8'],
                ['wvText', file_path, '-', '--charset=windows-1255'],
                ['wvText', file_path, '-', '--charset=hebrew'],
                ['wvText', file_path, '-', '--charset=cp1255'],
                ['wvText', file_path, '-', '--charset=iso-8859-8'],
                ['wvText', file_path, '-', '--charset=utf-8', '--config-file=/usr/local/share/wv/wvText.xml'],
                ['wvText', file_path, '-', '--charset=windows-1255', '--config-file=/usr/local/share/wv/wvText.xml']
            ]
            
            for cmd in wv_options:
                try:
                    result = subprocess.run(cmd, capture_output=True, text=True, check=True)
                    if result.stdout.strip():
                        for encoding in encodings:
                            try:
                                text = result.stdout.encode().decode(encoding)
                                if text and any(c.isalnum() for c in text):
                                    return text
                            except Exception:
                                continue
                except Exception as e:
                    print(f"wv attempt failed with command {' '.join(cmd)}: {str(e)}")
                    continue

            # Try olefile
            try:
                ole = olefile.OleFileIO(file_path)
                streams = ole.listdir()
                
                if b'WordDocument' in streams:
                    with ole.openstream('WordDocument') as stream:
                        content = stream.read()
                        for encoding in encodings:
                            try:
                                text = content.decode(encoding)
                                if text and any(c.isalnum() for c in text):
                                    return text
                            except Exception:
                                continue
            except Exception as e:
                print(f"olefile attempt failed: {str(e)}")

            # Try antiword with different encodings
            antiword_options = [
                ['antiword', file_path],
                ['antiword', '-m', 'UTF-8.txt', file_path],
                ['antiword', '-m', 'windows-1255.txt', file_path],
                ['antiword', '-m', 'hebrew.txt', file_path],
                ['antiword', '-m', 'cp1255.txt', file_path],
                ['antiword', '-m', 'iso-8859-8.txt', file_path]
            ]
            
            for cmd in antiword_options:
                try:
                    result = subprocess.run(cmd, capture_output=True, text=True, check=True)
                    if result.stdout.strip():
                        for encoding in encodings:
                            try:
                                text = result.stdout.encode().decode(encoding)
                                if text and any(c.isalnum() for c in text):
                                    return text
                            except Exception:
                                continue
                except Exception as e:
                    print(f"antiword attempt failed with command {' '.join(cmd)}: {str(e)}")
                    continue

            # Try python-docx as last resort
            try:
                doc = docx.Document(file_path)
                text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                if text and any(c.isalnum() for c in text):
                    return text
            except Exception as e:
                print(f"python-docx attempt failed: {str(e)}")

            # If all methods fail, try msoffcrypto
            try:
                with open(file_path, 'rb') as f:
                    file = msoffcrypto.OfficeFile(f)
                    file.load_key(password='VelvetSweatshop')  # Default password for older Office files
                    
                    with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                        file.decrypt(temp_file)
                        temp_file_path = temp_file.name
                    
                    # Try reading the decrypted file with python-docx
                    try:
                        doc = docx.Document(temp_file_path)
                        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                        os.unlink(temp_file_path)
                        if text and any(c.isalnum() for c in text):
                            return text
                    except Exception:
                        os.unlink(temp_file_path)
            except Exception as e:
                print(f"msoffcrypto attempt failed: {str(e)}")

        except Exception as e:
            print(f"Error extracting text from {file_path}: {str(e)}")
        
        return ""  # Return empty string if all methods fail

    def _process_doc_file(self, file_path: str) -> str:
        """Process DOC file and return its text content."""
        try:
            # Use antiword to extract text from DOC file
            try:
                import subprocess
                result = subprocess.run(['antiword', file_path], capture_output=True, text=True)
                if result.returncode == 0:
                    return result.stdout
                else:
                    self.logger.error(f"antiword failed with return code {result.returncode}")
                    return ""
            except Exception as e:
                self.logger.error(f"Error running antiword on {file_path}: {str(e)}")
                return ""
                
        except Exception as e:
            self.logger.error(f"Error processing DOC file {file_path}: {str(e)}")
            return ""

    def _process_docx_file(self, file_path: str) -> str:
        """Process DOCX file and return its text content."""
        try:
            import docx2txt
            try:
                return docx2txt.process(file_path)
            except Exception:
                self.logger.debug(f"docx2txt failed for {file_path}")
                
            # Try python-docx as fallback
            try:
                from docx import Document
                doc = Document(file_path)
                return "\n".join([paragraph.text for paragraph in doc.paragraphs])
            except Exception:
                self.logger.debug(f"python-docx failed for {file_path}")
                
            raise Exception("All DOCX processing methods failed")
                
        except Exception as e:
            self.logger.error(f"Error processing DOCX file {file_path}: {str(e)}")
            return ""

    @staticmethod
    def _clean_phone(phone: str) -> str:
        if not phone:
            return ""
        # Remove all non-digit characters except + for international numbers
        phone = "".join(c for c in phone if c.isdigit() or c == '+')
        return phone.strip()

    @staticmethod
    def _clean_email(email: str) -> str:
        if not email:
            return ""
        # Basic email cleaning
        return email.strip().lower()

    @staticmethod
    def _clean_address(address: str) -> str:
        if not address:
            return ""
        # Remove extra spaces and normalize
        address = " ".join(address.split())
        return address.strip()

    def add_phone(self, phone: str) -> bool:
        clean_phone = self._clean_phone(phone)
        if clean_phone and clean_phone not in self.phones:
            self.phones.append(clean_phone)
            return True
        return False

    def add_email(self, email: str) -> bool:
        clean_email = self._clean_email(email)
        if clean_email and clean_email not in self.emails:
            self.emails.append(clean_email)
            return True
        return False

    def add_address(self, address: str) -> bool:
        clean_address = self._clean_address(address)
        if clean_address and clean_address not in self.addresses:
            self.addresses.append(clean_address)
            return True
        return False

    def add_role(self, role: str) -> bool:
        if role and role.strip():
            self.role = role.strip()
            return True
        return False

    def merge(self, other: 'Contact') -> None:
        """
        Merge another contact's information into this contact
        """
        if other.name:
            self.name = other.name
        for phone in other.phones:
            self.add_phone(phone)
        for email in other.emails:
            self.add_email(email)
        for address in other.addresses:
            self.add_address(address)
        if other.role:
            self.add_role(other.role)
        if other.source_file:
            self.source_file = other.source_file

    def is_valid(self) -> bool:
        """
        Check if the contact has at least a name and either a phone number or email
        """
        return bool(self.name and (self.phones or self.emails))

    def to_dict(self) -> Dict[str, str]:
        """
        Convert contact to dictionary format
        """
        return {
            'name': self.name or '',
            'phone': self.phones[0] if self.phones else '',
            'email': self.emails[0] if self.emails else '',
            'address': self.addresses[0] if self.addresses else '',
            'source_file': self.source_file or ''
        }

    def extract_from_xlsx(self, file_path: str) -> Dict[str, List[str]]:
        """Extract contacts from Excel file."""
        try:
            names = []
            phones = []
            emails = []
            addresses = []
            
            # Define possible Hebrew column headers
            name_headers = ['שם', 'שם מלא', 'שם פרטי ומשפחה', 'איש קשר', 'נציג', 'שם הלקוח', 'לקוח']
            phone_headers = ['טלפון', 'נייד', 'טל', 'פלאפון', 'מספר טלפון', 'טלפון נייד', 'סלולרי']
            email_headers = ['מייל', 'אימייל', 'דוא"ל', 'דואל', 'כתובת מייל', 'EMAIL', 'E-MAIL']
            address_headers = ['כתובת', 'עיר', 'ישוב', 'מיקום', 'כתובת מלאה', 'רחוב']

            # Try pandas first
            try:
                df = pd.read_excel(file_path)
                
                # Convert all column names to string and lowercase for comparison
                df.columns = df.columns.map(str).str.strip().str.lower()
                
                # Find relevant columns
                name_col = None
                phone_col = None
                email_col = None
                address_col = None
                
                for col in df.columns:
                    col_lower = col.lower()
                    # Check for name column
                    if any(header.lower() in col_lower for header in name_headers):
                        name_col = col
                    # Check for phone column
                    elif any(header.lower() in col_lower for header in phone_headers):
                        phone_col = col
                    # Check for email column
                    elif any(header.lower() in col_lower for header in email_headers):
                        email_col = col
                    # Check for address column
                    elif any(header.lower() in col_lower for header in address_headers):
                        address_col = col
                
                # Process each row
                for _, row in df.iterrows():
                    # Extract name
                    if name_col and pd.notna(row[name_col]):
                        name = str(row[name_col]).strip()
                        if name and any(c.isalpha() for c in name):
                            names.append(name)
                    
                    # Extract phone
                    if phone_col and pd.notna(row[phone_col]):
                        phone = str(row[phone_col]).strip()
                        extracted_phones = self._extract_phones(phone)
                        phones.extend(extracted_phones)
                    
                    # Extract email
                    if email_col and pd.notna(row[email_col]):
                        email = str(row[email_col]).strip()
                        extracted_emails = self._extract_emails(email)
                        emails.extend(extracted_emails)
                    
                    # Extract address
                    if address_col and pd.notna(row[address_col]):
                        address = str(row[address_col]).strip()
                        if address and any(c.isalpha() for c in address):
                            addresses.append(address)
                    
                    # Also check entire row text for additional information
                    row_text = ' '.join(str(val) for val in row.values if pd.notna(val))
                    phones.extend(self._extract_phones(row_text))
                    emails.extend(self._extract_emails(row_text))
                
            except Exception as e:
                self.logger.debug(f"Pandas failed: {str(e)}")
                
                # Fallback to openpyxl
                workbook = openpyxl.load_workbook(file_path, data_only=True)
                
                for sheet in workbook.worksheets:
                    # Get headers from first row
                    headers = [str(cell.value).strip().lower() if cell.value else '' for cell in sheet[1]]
                    
                    # Find relevant columns
                    name_col = -1
                    phone_col = -1
                    email_col = -1
                    address_col = -1
                    
                    for idx, header in enumerate(headers):
                        # Check for name column
                        if any(h.lower() in header for h in name_headers):
                            name_col = idx
                        # Check for phone column
                        elif any(h.lower() in header for h in phone_headers):
                            phone_col = idx
                        # Check for email column
                        elif any(h.lower() in header for h in email_headers):
                            email_col = idx
                        # Check for address column
                        elif any(h.lower() in header for h in address_headers):
                            address_col = idx
                    
                    # Process each row
                    for row in list(sheet.rows)[1:]:  # Skip header row
                        # Extract name
                        if name_col >= 0 and row[name_col].value:
                            name = str(row[name_col].value).strip()
                            if name and any(c.isalpha() for c in name):
                                names.append(name)
                        
                        # Extract phone
                        if phone_col >= 0 and row[phone_col].value:
                            phone = str(row[phone_col].value).strip()
                            extracted_phones = self._extract_phones(phone)
                            phones.extend(extracted_phones)
                        
                        # Extract email
                        if email_col >= 0 and row[email_col].value:
                            email = str(row[email_col].value).strip()
                            extracted_emails = self._extract_emails(email)
                            emails.extend(extracted_emails)
                        
                        # Extract address
                        if address_col >= 0 and row[address_col].value:
                            address = str(row[address_col].value).strip()
                            if address and any(c.isalpha() for c in address):
                                addresses.append(address)
                        
                        # Also check entire row text for additional information
                        row_text = ' '.join(str(cell.value) for cell in row if cell.value)
                        phones.extend(self._extract_phones(row_text))
                        emails.extend(self._extract_emails(row_text))
            
            # Remove duplicates while preserving order
            names = list(dict.fromkeys(names))
            phones = list(dict.fromkeys(phones))
            emails = list(dict.fromkeys(emails))
            addresses = list(dict.fromkeys(addresses))
            
            return {
                "names": names,
                "phones": phones,
                "emails": emails,
                "addresses": addresses
            }

        except Exception as e:
            self.logger.error(f"Error processing Excel file {file_path}: {str(e)}")
            return {
                "names": [],
                "phones": [],
                "emails": [],
                "addresses": []
            }

    def extract_from_doc(self, file_path: str) -> Dict[str, List[str]]:
        """Extract contacts from DOC/DOCX file."""
        try:
            names = []
            phones = []
            emails = []
            addresses = []
            text = ""
            
            # For DOCX files
            if file_path.lower().endswith('.docx'):
                try:
                    # Try python-docx first
                    doc = Document(file_path)
                    
                    # Process tables
                    for table in doc.tables:
                        # Get headers from first row
                        headers = []
                        if table.rows:
                            headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
                        
                        # Find relevant columns
                        name_col = -1
                        phone_col = -1
                        email_col = -1
                        address_col = -1
                        
                        for idx, header in enumerate(headers):
                            if any(word in header for word in ['שם', 'איש קשר', 'נציג']):
                                name_col = idx
                            elif any(word in header for word in ['טלפון', 'נייד', 'טל', 'פלאפון']):
                                phone_col = idx
                            elif any(word in header for word in ['מייל', 'אימייל', 'דוא"ל', 'דואל']):
                                email_col = idx
                            elif any(word in header for word in ['כתובת', 'עיר', 'ישוב']):
                                address_col = idx
                        
                        # Process each row
                        for row in table.rows[1:]:  # Skip header row
                            row_cells = row.cells
                            row_text = ' '.join(cell.text.strip() for cell in row_cells if cell.text.strip())
                            
                            # Extract from specific columns if found
                            if name_col >= 0 and len(row_cells) > name_col:
                                name = row_cells[name_col].text.strip()
                                if name and any(c.isalpha() for c in name):
                                    names.append(name)
                            
                            if phone_col >= 0 and len(row_cells) > phone_col:
                                phones.extend(self._extract_phones(row_cells[phone_col].text))
                            
                            if email_col >= 0 and len(row_cells) > email_col:
                                emails.extend(self._extract_emails(row_cells[email_col].text))
                            
                            if address_col >= 0 and len(row_cells) > address_col:
                                address = row_cells[address_col].text.strip()
                                if address and any(c.isalpha() for c in address):
                                    addresses.append(address)
                            
                            # Also check entire row text
                            phones.extend(self._extract_phones(row_text))
                            emails.extend(self._extract_emails(row_text))
                    
                    # Process paragraphs
                    for para in doc.paragraphs:
                        text += para.text + "\n"
                        
                except Exception as e:
                    self.logger.debug(f"python-docx failed: {str(e)}")
                    # Fallback to docx2txt
                    try:
                        text = docx2txt.process(file_path)
                    except Exception as e:
                        self.logger.debug(f"docx2txt failed: {str(e)}")
            
            # For DOC files
            else:
                # Try multiple methods to extract text from DOC files
                text = self._extract_text_from_old_doc(file_path)
            
            # Process extracted text
            if text:
                # Split text into lines
                lines = text.split('\n')
                
                # Process each line
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    # Try to extract name and role
                    name_and_role = self._extract_name_and_role(line)
                    if name_and_role:
                        name, _ = name_and_role
                        if name:
                            names.append(name)
                    
                    # Extract other information
                    phones.extend(self._extract_phones(line))
                    emails.extend(self._extract_emails(line))
                    
                    # Check for address patterns
                    if any(word in line.lower() for word in ['רחוב', 'שדרות', 'כתובת:', 'עיר:', 'ישוב:']):
                        addresses.append(line)
            
            # Remove duplicates while preserving order
            names = list(dict.fromkeys(names))
            phones = list(dict.fromkeys(phones))
            emails = list(dict.fromkeys(emails))
            addresses = list(dict.fromkeys(addresses))
            
            return {
                "names": names,
                "phones": phones,
                "emails": emails,
                "addresses": addresses
            }

        except Exception as e:
            self.logger.error(f"Error extracting from DOC/DOCX {file_path}: {str(e)}")
            return {
                "names": [],
                "phones": [],
                "emails": [],
                "addresses": []
            }

    def extract_from_pdf(self, file_path: str) -> Dict[str, List[str]]:
        """Extract contacts from PDF file."""
        try:
            names = []
            phones = []
            emails = []
            addresses = []
            text = ""
            
            # Try pdfplumber first
            try:
                with pdfplumber.open(file_path) as pdf:
                    # Process each page
                    for page in pdf.pages:
                        # Extract text
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                        
                        # Extract tables
                        tables = page.extract_tables()
                        for table in tables:
                            # Get headers from first row
                            headers = [str(cell).strip().lower() if cell else '' for cell in table[0]]
                            
                            # Find relevant columns
                            name_col = -1
                            phone_col = -1
                            email_col = -1
                            address_col = -1
                            
                            for idx, header in enumerate(headers):
                                if any(word in header for word in ['שם', 'איש קשר', 'נציג']):
                                    name_col = idx
                                elif any(word in header for word in ['טלפון', 'נייד', 'טל', 'פלאפון']):
                                    phone_col = idx
                                elif any(word in header for word in ['מייל', 'אימייל', 'דוא"ל', 'דואל']):
                                    email_col = idx
                                elif any(word in header for word in ['כתובת', 'עיר', 'ישוב']):
                                    address_col = idx
                            
                            # Process each row
                            for row in table[1:]:  # Skip header row
                                row_text = ' '.join(str(cell).strip() for cell in row if cell)
                                
                                # Extract from specific columns if found
                                if name_col >= 0 and len(row) > name_col and row[name_col]:
                                    name = str(row[name_col]).strip()
                                    if name and any(c.isalpha() for c in name):
                                        names.append(name)
                                
                                if phone_col >= 0 and len(row) > phone_col and row[phone_col]:
                                    phones.extend(self._extract_phones(str(row[phone_col])))
                                
                                if email_col >= 0 and len(row) > email_col and row[email_col]:
                                    emails.extend(self._extract_emails(str(row[email_col])))
                                
                                if address_col >= 0 and len(row) > address_col and row[address_col]:
                                    address = str(row[address_col]).strip()
                                    if address and any(c.isalpha() for c in address):
                                        addresses.append(address)
                                
                                # Also check entire row text
                                phones.extend(self._extract_phones(row_text))
                                emails.extend(self._extract_emails(row_text))
            
            except Exception as e:
                self.logger.debug(f"pdfplumber failed: {str(e)}")
                # Fallback to PyPDF2
                try:
                    with open(file_path, 'rb') as file:
                        reader = PyPDF2.PdfReader(file)
                        for page in reader.pages:
                            text += page.extract_text() + "\n"
                except Exception as e:
                    self.logger.debug(f"PyPDF2 failed: {str(e)}")
                    
                    # If both methods fail, try OCR
                    try:
                        # Convert PDF to images
                        images = convert_from_path(file_path)
                        
                        # Process each page image
                        for image in images:
                            # Perform OCR with Hebrew support
                            text += pytesseract.image_to_string(image, lang='heb+eng') + "\n"
                    except Exception as e:
                        self.logger.debug(f"OCR failed: {str(e)}")
            
            # Process extracted text
            if text:
                # Split text into lines
                lines = text.split('\n')
                
                # Process each line
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    # Try to extract name and role
                    name_and_role = self._extract_name_and_role(line)
                    if name_and_role:
                        name, _ = name_and_role
                        if name:
                            names.append(name)
                    
                    # Extract other information
                    phones.extend(self._extract_phones(line))
                    emails.extend(self._extract_emails(line))
                    
                    # Check for address patterns
                    if any(word in line.lower() for word in ['רחוב', 'שדרות', 'כתובת:', 'עיר:', 'ישוב:']):
                        addresses.append(line)
            
            # Remove duplicates while preserving order
            names = list(dict.fromkeys(names))
            phones = list(dict.fromkeys(phones))
            emails = list(dict.fromkeys(emails))
            addresses = list(dict.fromkeys(addresses))
            
            return {
                "names": names,
                "phones": phones,
                "emails": emails,
                "addresses": addresses
            }

        except Exception as e:
            self.logger.error(f"Error extracting from PDF {file_path}: {str(e)}")
            return {
                "names": [],
                "phones": [],
                "emails": [],
                "addresses": []
            }

    def _extract_phones(self, text: str) -> List[str]:
        """Extract phone numbers from text."""
        phones = []
        matches = re.finditer(self.phone_pattern, text)
        for match in matches:
            phone = match.group()
            # Clean up the phone number
            phone = re.sub(r'[^\d+]', '', phone)
            # Add leading zero if needed
            if len(phone) == 9 and not phone.startswith('0'):
                phone = '0' + phone
            # Add leading zero if needed for 8-digit numbers
            if len(phone) == 8 and not phone.startswith('0'):
                phone = '0' + phone
            # Normalize to international format if needed
            if phone.startswith('0'):
                phone = '+972' + phone[1:]
            if len(phone) >= 12:  # Only add valid phone numbers
                phones.append(phone)
        return list(set(phones))  # Remove duplicates

    def _extract_emails(self, text: str) -> List[str]:
        """Extract email addresses from text."""
        emails = set()
        
        # Basic email pattern
        email_pattern = r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}'
        
        # Find all matches
        matches = re.finditer(email_pattern, text)
        for match in matches:
            email = match.group().lower()
            # Additional validation
            if self._is_valid_email(email):
                emails.add(email)
                
        # Look for emails with Hebrew text around them
        hebrew_context_pattern = r'[\u0590-\u05FF\s]+({})[\u0590-\u05FF\s]+'.format(email_pattern)
        matches = re.finditer(hebrew_context_pattern, text)
        for match in matches:
            email = match.group(1).lower()
            if self._is_valid_email(email):
                emails.add(email)
        
        return sorted(list(emails))
    
    def _is_valid_email(self, email: str) -> bool:
        """Validate email address."""
        # Basic validation
        if not re.match(r'^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}$', email):
            return False
            
        # Check for common invalid patterns
        invalid_patterns = [
            r'example\.com$',
            r'test\.com$',
            r'domain\.com$',
            r'@.*@',
            r'\.{2,}',
            r'^[0-9]+@'
        ]
        
        for pattern in invalid_patterns:
            if re.search(pattern, email):
                return False
                
        # Check for valid domains
        domain = email.split('@')[1]
        valid_domains = [
            'gmail.com', 'yahoo.com', 'hotmail.com', 'walla.co.il',
            'gmail.co.il', 'bezeqint.net', 'zahav.net.il', 'netvision.net.il',
            'outlook.com', 'aol.com', 'mail.com', 'office.com'
        ]
        
        if '.' not in domain:
            return False
            
        return True

    def _extract_addresses(self, text: str) -> List[str]:
        """Extract addresses from text."""
        addresses = []
        matches = re.finditer(self.address_pattern, text)
        for match in matches:
            address = match.group().strip()
            addresses.append(address)
        return list(set(addresses))  # Remove duplicates

    def _extract_text_from_table(self, table) -> str:
        """Extract text from a table object."""
        text = []
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    text.append(cell.text.strip())
        return ' '.join(text)

    def load_existing_contacts(self):
        """Load existing contacts from the JSON file."""
        try:
            if os.path.exists('contacts.json'):
                with open('contacts.json', 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.contacts = [Contact(**contact) for contact in data]
                    self.logger.info(f"Loaded {len(self.contacts)} existing contacts")
        except Exception as e:
            self.logger.error(f"שגיאה בטעינת קובץ קיים: {str(e)}")

    def process_files(self):
        """Process all files in the directory."""
        try:
            # Load existing contacts first
            self.load_existing_contacts()

            # Get list of files to process
            files = []
            for root, _, filenames in os.walk('דוגמאות'):
                for filename in filenames:
                    if filename.endswith(('.xlsx', '.xls', '.doc', '.docx', '.pdf')):
                        files.append(os.path.join(root, filename))

            self.logger.info(f"נמצאו {len(files)} קבצים לעיבוד")

            # Process each file
            for i, file_path in enumerate(files, 1):
                try:
                    if file_path in self.processed_files:
                        continue

                    contacts_before = len(self.contacts)
                    
                    if file_path.endswith(('.xlsx', '.xls')):
                        self.extract_from_xlsx(file_path)
                    elif file_path.endswith(('.doc', '.docx')):
                        self.extract_from_doc(file_path)
                    elif file_path.endswith('.pdf'):
                        self.extract_from_pdf(file_path)

                    contacts_after = len(self.contacts)
                    new_contacts = contacts_after - contacts_before
                    if new_contacts > 0:
                        self.logger.info(f"נמצאו {new_contacts} אנשי קשר חדשים בקובץ {file_path}")

                    self.processed_files.add(file_path)

                except Exception as e:
                    self.logger.error(f"Error processing file {file_path}: {str(e)}")

                # Save progress every 5 files
                if i % 5 == 0:
                    self.save_contacts()
                    self.logger.info(f"נשמרה התקדמות אחרי {i} קבצים")

            # Final save
            self.save_contacts()
            self.logger.info(f"הסתיים עיבוד {len(files)} קבצים")
            self.logger.info(f"נמצאו {len(self.contacts)} אנשי קשר חדשים")
            self.logger.info(f"סה\"כ {len(self.contacts)} אנשי קשר נשמרו בקובץ")

        except Exception as e:
            self.logger.error(f"שגיאה בעיבוד הקבצים: {str(e)}")

    def save_contacts(self):
        """Save contacts to a JSON file."""
        if not self.contacts:
            print("אין אנשי קשר לשמירה")
            return

        try:
            contacts_data = [contact.__dict__ for contact in self.contacts]
            with open('contacts.json', 'w', encoding='utf-8') as f:
                json.dump(contacts_data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            self.logger.error(f"שגיאה בשמירת אנשי קשר: {str(e)}")

    def extract_from_text(self, text: str) -> Dict[str, List[str]]:
        """Extract contact information from text."""
        names = []
        phones = []
        emails = []
        
        # Extract names (sequences of Hebrew letters)
        for match in self.name_pattern.finditer(text):
            name = match.group().strip()
            if len(name) > 1 and not any(char.isdigit() for char in name):
                names.append(name)
        
        # Extract phone numbers
        for match in self.phone_pattern.finditer(text):
            phone = match.group().strip()
            # Clean up phone number
            phone = re.sub(r'[\s\-]', '', phone)
            if phone.startswith('972'):
                phone = '+' + phone
            elif phone.startswith('05'):
                phone = '+972' + phone[1:]
            phones.append(phone)
        
        # Extract email addresses
        for match in self.email_pattern.finditer(text):
            email = match.group().strip().lower()
            emails.append(email)
            
        return {
            'names': names,
            'phones': phones,
            'emails': emails
        }

class Contact:
    def __init__(self, name: str = None, phone: str = None, email: str = None, address: str = None, role: str = None, source_file: str = None):
        self.name = name
        self.phones: Set[str] = {phone} if phone else set()
        self.emails: Set[str] = {email} if email else set()
        self.addresses: Set[str] = {address} if address else set()
        self.roles: Set[str] = {role} if role else set()
        self.source_file = source_file

    def add_phone(self, phone: str) -> None:
        if phone:
            self.phones.add(phone)

    def add_email(self, email: str) -> None:
        if email:
            self.emails.add(email)

    def add_address(self, address: str) -> None:
        if address:
            self.addresses.add(address)

    def add_role(self, role: str) -> None:
        if role:
            self.roles.add(role)

    def merge(self, other: 'Contact') -> None:
        if other.name:
            self.name = other.name
        self.phones.update(other.phones)
        self.emails.update(other.emails)
        self.addresses.update(other.addresses)
        self.roles.update(other.roles)
        if other.source_file:
            self.source_file = other.source_file

    def is_valid(self) -> bool:
        return bool(self.name and (self.phones or self.emails))

    def to_dict(self) -> Dict[str, str]:
        return {
            'name': self.name or '',
            'phone': next(iter(self.phones), ''),
            'email': next(iter(self.emails), ''),
            'address': next(iter(self.addresses), ''),
            'role': next(iter(self.roles), ''),
            'additional_phones': ', '.join(list(self.phones)[1:]) if len(self.phones) > 1 else '',
            'additional_emails': ', '.join(list(self.emails)[1:]) if len(self.emails) > 1 else '',
            'additional_addresses': ', '.join(list(self.addresses)[1:]) if len(self.addresses) > 1 else '',
            'additional_roles': ', '.join(list(self.roles)[1:]) if len(self.roles) > 1 else '',
            'source_file': self.source_file or ''
        }

def organize_contacts(names: List[str], phones: List[str], emails: List[str], addresses: List[str]) -> List[Contact]:
    """Organize extracted information into Contact objects."""
    contacts = []
    
    # If we have names, create a contact for each name
    if names:
        for name in names:
            contact = Contact(name=name)
            contacts.append(contact)
    
    # If we have no names but have other information
    if not names and (phones or emails or addresses):
        contact = Contact()
        contacts.append(contact)
    
    # Add phones to contacts
    for phone in phones:
        if contacts:
            # Try to find a contact without a phone
            contact = next((c for c in contacts if not c.phones), contacts[0])
            contact.add_phone(phone)
        else:
            contact = Contact()
            contact.add_phone(phone)
            contacts.append(contact)
    
    # Add emails to contacts
    for email in emails:
        if contacts:
            # Try to find a contact without an email
            contact = next((c for c in contacts if not c.emails), contacts[0])
            contact.add_email(email)
        else:
            contact = Contact()
            contact.add_email(email)
            contacts.append(contact)
    
    # Add addresses to contacts
    for address in addresses:
        if contacts:
            # Try to find a contact without an address
            contact = next((c for c in contacts if not c.addresses), contacts[0])
            contact.add_address(address)
        else:
            contact = Contact()
            contact.add_address(address)
            contacts.append(contact)
    
    return contacts

def main():
    # Configure logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s'
    )

    # Initialize contact extractor
    extractor = ContactExtractor()
    
    # Dictionary to store all contacts
    all_contacts: Dict[str, Contact] = {}
    
    # Try to load existing contacts
    output_file = "UP Data Base.xlsx"
    if os.path.exists(output_file):
        try:
            existing_df = pd.read_excel(output_file)
            # Convert Hebrew column names back to English
            reverse_column_mapping = {
                'שם': 'name',
                'תפקיד': 'role',
                'טלפון': 'phone',
                'אימייל': 'email',
                'כתובת': 'address',
                'טלפונים נוספים': 'additional_phones',
                'אימיילים נוספים': 'additional_emails',
                'כתובות נוספות': 'additional_addresses',
                'תפקידים נוספים': 'additional_roles',
                'מקור': 'source_file'
            }
            existing_df = existing_df.rename(columns={v: k for k, v in reverse_column_mapping.items()})
            
            for _, row in existing_df.iterrows():
                contact = Contact(row['name'])
                if pd.notna(row['phone']): contact.add_phone(str(row['phone']))
                if pd.notna(row['email']): contact.add_email(row['email'])
                if pd.notna(row['address']): contact.add_address(row['address'])
                if pd.notna(row['role']): contact.add_role(row['role'])
                
                # Add additional information
                if pd.notna(row['additional_phones']):
                    for phone in str(row['additional_phones']).split(','):
                        contact.add_phone(phone.strip())
                if pd.notna(row['additional_emails']):
                    for email in row['additional_emails'].split(','):
                        contact.add_email(email.strip())
                if pd.notna(row['additional_addresses']):
                    for addr in row['additional_addresses'].split(','):
                        contact.add_address(addr.strip())
                if pd.notna(row['additional_roles']):
                    for role in row['additional_roles'].split(','):
                        contact.add_role(role.strip())
                
                if contact.is_valid():
                    all_contacts[contact.name] = contact
            
            logging.info(f"נטענו {len(all_contacts)} אנשי קשר קיימים")
        except Exception as e:
            logging.error(f"שגיאה בטעינת קובץ קיים: {str(e)}")

    # Get list of files to process
    doc_dir = '../מתווכים'
    files = []
    for root, _, filenames in os.walk(doc_dir):
        for filename in filenames:
            if filename.endswith(('.xlsx', '.xls', '.docx', '.doc', '.pdf', '.txt')):
                files.append(os.path.join(root, filename))

    logging.info(f"נמצאו {len(files)} קבצים לעיבוד")
    
    # Process each file
    files_processed = 0
    contacts_found = 0
    
    for file_path in files:
        try:
            # Process the file
            result = process_file(extractor, file_path)
            
            if result:
                # Organize contacts from the results
                names = result.get('names', [])
                phones = result.get('phones', [])
                emails = result.get('emails', [])
                addresses = result.get('addresses', [])
                
                new_contacts = organize_contacts(names, phones, emails, addresses)
                
                # Add source file information
                for contact in new_contacts:
                    contact.source_file = os.path.basename(file_path)
                    
                    # Add to all_contacts if valid
                    if contact.is_valid():
                        if contact.name in all_contacts:
                            all_contacts[contact.name].merge(contact)
                        else:
                            all_contacts[contact.name] = contact
                            contacts_found += 1
            
            files_processed += 1
            
            # Save progress every 5 files
            if files_processed % 5 == 0:
                save_contacts_to_excel(all_contacts, output_file)
                logging.info(f"נשמרה התקדמות אחרי {files_processed} קבצים")
                
        except Exception as e:
            print(f"שגיאה בעיבוד המסמך {os.path.basename(file_path)}: {str(e)}")
            continue
    
    # Save final results
    save_contacts_to_excel(all_contacts, output_file)
    
    logging.info(f"הסתיים עיבוד {files_processed} קבצים")
    logging.info(f"נמצאו {contacts_found} אנשי קשר חדשים")
    logging.info(f"סה\"כ {len(all_contacts)} אנשי קשר נשמרו בקובץ")

def process_file(extractor: ContactExtractor, file_path: str) -> Optional[Dict[str, List[str]]]:
    """Process a single file and return extracted contacts."""
    try:
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension in ['.xlsx', '.xls']:
            extracted = process_excel_file(file_path)
        elif file_extension in ['.doc', '.docx']:
            if file_extension == '.doc':
                text = process_doc_file(file_path)
            else:
                text = process_docx_file(file_path)
            extracted = extractor.extract_from_text(text)
        elif file_extension == '.pdf':
            text = process_pdf_file(file_path)
            extracted = extractor.extract_from_text(text)
        else:
            logging.warning(f"סוג קובץ לא נתמך: {file_extension}")
            return None
            
        # Clean up extracted data
        if extracted:
            extracted['phones'] = [extractor._clean_phone(p) for p in extracted.get('phones', [])]
            extracted['emails'] = [extractor._clean_email(e) for e in extracted.get('emails', [])]
            extracted['addresses'] = [extractor._clean_address(a) for a in extracted.get('addresses', [])]
            extracted['names'] = [extractor._normalize_name(n) for n in extracted.get('names', [])]
            
        return extracted
    except Exception as e:
        logging.error(f"שגיאה בעיבוד קובץ {file_path}: {str(e)}")
        return None

def process_docx_file(file_path: str) -> str:
    """Process a DOCX file and extract text."""
    try:
        doc = Document(file_path)
        paragraphs_text = [paragraph.text for paragraph in doc.paragraphs]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                tables_text.append(" ".join(cell.text for cell in row.cells))
        return "\n".join(paragraphs_text + tables_text)
    except Exception as e:
        logging.error(f"שגיאה בעיבוד קובץ DOCX {file_path}: {str(e)}")
        return ""

def save_contacts_to_excel(contacts: Dict[str, Contact], output_file: str) -> None:
    """Save contacts to Excel file."""
    try:
        # Convert contacts to list of dictionaries
        contacts_data = [contact.to_dict() for contact in contacts.values()]
        
        # Create DataFrame
        df = pd.DataFrame(contacts_data)
        
        # Define column order and Hebrew names
        column_mapping = {
            'name': 'שם',
            'role': 'תפקיד',
            'phone': 'טלפון',
            'email': 'אימייל',
            'address': 'כתובת',
            'additional_phones': 'טלפונים נוספים',
            'additional_emails': 'אימיילים נוספים',
            'additional_addresses': 'כתובות נוספות',
            'additional_roles': 'תפקידים נוספים',
            'source_file': 'מקור'
        }
        
        # Reorder and rename columns
        df = df.reindex(columns=list(column_mapping.keys()))
        df = df.rename(columns=column_mapping)
        
        # Create Excel writer
        with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
            # Write DataFrame to Excel
            df.to_excel(writer, index=False, sheet_name='אנשי קשר')
            
            # Get the worksheet
            worksheet = writer.sheets['אנשי קשר']
            
            # Set column widths
            for idx, col in enumerate(df.columns):
                max_length = max(
                    df[col].astype(str).apply(len).max(),  # Length of largest item
                    len(str(col))  # Length of column name
                ) + 2  # Adding a little extra space
                worksheet.column_dimensions[get_column_letter(idx + 1)].width = max_length
            
            # Set RTL direction and alignment
            for row in worksheet.iter_rows():
                for cell in row:
                    cell.alignment = Alignment(horizontal='right', vertical='center')
            
            worksheet.sheet_view.rightToLeft = True
        
        logging.info(f"נשמרו {len(contacts)} אנשי קשר לקובץ {output_file}")
    
    except Exception as e:
        logging.error(f"שגיאה בשמירת אנשי קשר לקובץ Excel: {str(e)}")

def download_tika_server(jar_path):
    url = "https://repo1.maven.org/maven2/org/apache/tika/tika-server-standard/2.6.0/tika-server-standard-2.6.0.jar"
    try:
        # Create a pool manager with certificate verification
        http = urllib3.PoolManager(
            cert_reqs='CERT_REQUIRED',
            ca_certs=certifi.where()
        )
        
        # Download the file
        with http.request('GET', url, preload_content=False) as resp, open(jar_path, 'wb') as out_file:
            if resp.status != 200:
                raise Exception(f"Failed to download file: HTTP {resp.status}")
            shutil.copyfileobj(resp, out_file)
            
        return True
    except Exception as e:
        logging.error(f"Failed to download Tika server: {str(e)}")
        return False

def process_doc_file(file_path):
    """Process a DOC file using multiple methods."""
    try:
        # First try docx2python
        try:
            from docx2python import docx2python
            doc = docx2python(file_path)
            text = doc.text
            if text.strip():
                return text
        except Exception as e:
            logging.debug(f"docx2python failed: {str(e)}")
        
        # Then try oletools
        try:
            from oletools.olevba import VBA_Parser
            vba = VBA_Parser(file_path)
            text = vba.get_macros()
            if text and any(t.strip() for t in text):
                return '\n'.join(text)
        except Exception as e:
            logging.debug(f"oletools failed: {str(e)}")
            
        # Try olefile
        try:
            import olefile
            ole = olefile.OleFileIO(file_path)
            if ole.exists('WordDocument'):
                with ole.openstream('WordDocument') as stream:
                    text = stream.read().decode('utf-16le', errors='ignore')
                    if text.strip():
                        return text
        except Exception as e:
            logging.debug(f"olefile failed: {str(e)}")
            
        # Try converting with LibreOffice
        try:
            output_path = file_path + ".docx"
            convert_cmd = f"soffice --headless --convert-to docx:\"MS Word 2007 XML\" --outdir {os.path.dirname(file_path)} {file_path}"
            subprocess.run(convert_cmd, shell=True, check=True, capture_output=True)
            if os.path.exists(output_path):
                doc = Document(output_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                os.remove(output_path)  # Clean up
                if text.strip():
                    return text
        except Exception as e:
            logging.debug(f"LibreOffice conversion failed: {str(e)}")
            
        # Finally try Tika
        try:
            from tika import parser
            parsed = parser.from_file(file_path)
            if parsed and parsed.get("content", "").strip():
                return parsed["content"]
        except Exception as e:
            logging.debug(f"Tika failed: {str(e)}")
            
        raise Exception("All DOC processing methods failed")
            
    except Exception as e:
        logging.error(f"All DOC processing methods failed for {file_path}: {str(e)}")
        return ""

def process_pdf_file(file_path):
    """Process a PDF file using multiple methods."""
    try:
        # First try pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                text = "\n".join(page.extract_text() for page in pdf.pages)
                if text.strip():
                    return text
        except Exception as e:
            logging.debug(f"pdfplumber failed: {str(e)}")
            
        # Then try PyPDF2
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            text = "\n".join(page.extract_text() for page in reader.pages)
            if text.strip():
                return text
        except Exception as e:
            logging.debug(f"PyPDF2 failed: {str(e)}")
            
        # Try OCR with pdf2image and pytesseract
        try:
            from pdf2image import convert_from_path
            import pytesseract
            
            # Convert PDF to images
            images = convert_from_path(file_path)
            
            # Extract text from each image
            text = "\n".join(pytesseract.image_to_string(image, lang='heb+eng') for image in images)
            if text.strip():
                return text
        except Exception as e:
            logging.debug(f"OCR failed: {str(e)}")
            
        # Finally try Tika
        try:
            from tika import parser
            parsed = parser.from_file(file_path)
            if parsed and parsed.get("content", "").strip():
                return parsed["content"]
        except Exception as e:
            logging.debug(f"Tika failed: {str(e)}")
            
        raise Exception("All PDF processing methods failed")
            
    except Exception as e:
        logging.error(f"All PDF processing methods failed for {file_path}: {str(e)}")
        return ""

def process_excel_file(file_path):
    """Process an Excel file and extract contact information."""
    try:
        import pandas as pd
        
        # Try reading with pandas
        try:
            # First try auto-detection of the engine  
            df = pd.read_excel(file_path)
        except Exception:
            # If that fails, try explicitly with openpyxl
            df = pd.read_excel(file_path, engine='openpyxl')
            
        # Convert all columns to string and handle NaN values
        df = df.astype(str).replace('nan', '')
        
        # Common Hebrew column names for contact information
        name_columns = ['שם', 'שם מלא', 'שם פרטי', 'שם משפחה', 'איש קשר', 'שם איש קשר']
        phone_columns = ['טלפון', 'נייד', 'טלפון נייד', 'מספר טלפון', 'טל', 'פלאפון', 'סלולרי']
        email_columns = ['אימייל', 'מייל', 'דוא"ל', 'דואר אלקטרוני', 'email', 'e-mail']
        address_columns = ['כתובת', 'רחוב', 'מען', 'עיר', 'ישוב']
        
        # Initialize lists for extracted information
        names = []
        phones = []
        emails = []
        addresses = []
        
        # Function to find matching columns
        def find_matching_columns(df, patterns):
            matches = []
            for col in df.columns:
                col_str = str(col).strip().lower()
                if any(pattern.lower() in col_str for pattern in patterns):
                    matches.append(col)
            return matches
        
        # Extract information from matching columns
        for col in find_matching_columns(df, name_columns):
            names.extend(df[col].dropna().astype(str).str.strip().tolist())
            
        for col in find_matching_columns(df, phone_columns):
            phones.extend(df[col].dropna().astype(str).str.strip().tolist())
            
        for col in find_matching_columns(df, email_columns):
            emails.extend(df[col].dropna().astype(str).str.strip().tolist())
            
        for col in find_matching_columns(df, address_columns):
            addresses.extend(df[col].dropna().astype(str).str.strip().tolist())
            
        # If no columns were found, try to extract from all columns
        if not any([names, phones, emails, addresses]):
            text = "\n".join(df.values.astype(str).flatten())
            return extract_from_text(text)
            
        return {
            'names': names,
            'phones': phones,
            'emails': emails,
            'addresses': addresses
        }
            
    except Exception as e:
        logging.error(f"Failed to process Excel file {file_path}: {str(e)}")
        return {
            'names': [],
            'phones': [],
            'emails': [],
            'addresses': []
        }

def clean_and_validate_data(extracted: Dict[str, List[str]]) -> Dict[str, List[str]]:
    """Clean and validate extracted data."""
    if not extracted:
        return {
            'names': [],
            'phones': [],
            'emails': [],
            'addresses': []
        }
        
    # Clean phone numbers
    phones = []
    for phone in extracted.get('phones', []):
        # Remove common prefixes
        phone = re.sub(r'^0', '', phone)
        phone = re.sub(r'^972', '', phone)
        phone = re.sub(r'^\+972', '', phone)
        
        # Remove non-digit characters
        phone = re.sub(r'\D', '', phone)
        
        # Add international prefix
        if len(phone) >= 9:
            phone = '+972' + phone[-9:]
            phones.append(phone)
            
    # Clean emails
    emails = []
    email_pattern = re.compile(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$')
    for email in extracted.get('emails', []):
        email = email.lower().strip()
        if email_pattern.match(email):
            emails.append(email)
            
    # Clean names
    names = []
    for name in extracted.get('names', []):
        # Remove multiple spaces
        name = re.sub(r'\s+', ' ', name.strip())
        
        # Remove digits and special characters
        name = re.sub(r'[0-9!"#$%&\'()*+,-./:;<=>?@[\]^_`{|}~]', '', name)
        
        if len(name) >= 2 and all(c.isalpha() or c.isspace() for c in name):
            names.append(name)
            
    # Clean addresses
    addresses = []
    for address in extracted.get('addresses', []):
        # Remove multiple spaces
        address = re.sub(r'\s+', ' ', address.strip())
        
        if len(address) >= 5:
            addresses.append(address)
            
    return {
        'names': names,
        'phones': phones,
        'emails': emails,
        'addresses': addresses
    }

def extract_from_text(text: str) -> Dict[str, List[str]]:
    """Extract contact information from text."""
    if not text or not isinstance(text, str):
        return {
            'names': [],
            'phones': [],
            'emails': [],
            'addresses': []
        }
        
    # Initialize patterns
    name_pattern = re.compile(r'[א-ת\s]{2,}')
    phone_pattern = re.compile(r'(?:\+972|05|972|0)[0-9\-\s]{8,}')
    email_pattern = re.compile(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}')
    address_pattern = re.compile(r'(?:רחוב|רח\'|שד\'|שדרות|דרך|כביש|סמטת|סמ\')\s+[\u0590-\u05FF\s\w.,\-\'\"]+')
    
    # Extract information
    names = [match.group().strip() for match in name_pattern.finditer(text)]
    phones = [match.group().strip() for match in phone_pattern.finditer(text)]
    emails = [match.group().strip().lower() for match in email_pattern.finditer(text)]
    addresses = [match.group().strip() for match in address_pattern.finditer(text)]
    
    # Clean and validate the extracted data
    return clean_and_validate_data({
        'names': names,
        'phones': phones,
        'emails': emails,
        'addresses': addresses
    })

if __name__ == '__main__':
    main() 