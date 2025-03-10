import os
from docx import Document
import subprocess
from pathlib import Path
import pandas as pd
import re
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
import docx2txt
import tempfile
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

def clean_name(name):
    if not name:
        return name
    
    # ניקוי השם
    name = str(name).strip()
    # הסרת תווי שורה חדשה
    name = re.sub(r'\s*\n\s*', ' ', name)
    
    # הסרת מילות סטטוס ומילים נפוצות
    status_words = [
        'בטיפול', 'במעקב', 'אחרי החגים', 'לא רלוונטי', 'הועבר ל', 'נרשם', 'נרשמה', 'מעוניין', 'לא מעוניין',
        'לא כרגע', 'לא עונה', 'ניסיתי', 'הועבר', 'דוח', 'מעקב', 'פרסום', 'מס', 'ות', 'ונטי', 'קורסים',
        'to study', 'study', 'פגישה', 'מחר', 'יום', 'אשתו', 'בעלה', 'זאוס', 'מוטי', 'טלפון', 'נייד',
        'שם', 'משפחה', 'כתובת', 'מייל', 'הערות', 'email', 'e-mail', 'מתעניין', 'סוכן', 'זכיין', 'רשימת',
        'משתתפים', 'תלמידים', 'פוטנציאליים', 'קורס', 'מבחן', 'הכנה', 'לבחינה'
    ]
    for word in status_words:
        name = re.sub(rf'\b{word}\b', '', name, flags=re.IGNORECASE)
    
    # הסרת סוגריים ותוכנם
    name = re.sub(r'\([^)]*\)', '', name)
    
    # הסרת שמות ערים ושכונות
    cities = [
        'תל אביב', 'רמת גן', 'חיפה', 'ירושלים', 'באר שבע', 'רחובות', 'פתח תקווה', 'רעננה', 'הרצליה',
        'גבעתיים', 'חולון', 'בת ים', 'נתניה', 'אשדוד', 'אשקלון', 'רמלה', 'לוד', 'כפר סבא', 'רמת השרון',
        'ראש העין', 'פ"ת', 'פ״ת', 'גבעת', 'ראש', 'ת"א', 'צור משה', 'כפר יונה', 'פרדס חנה', 'מודיעין',
        'גני תקווה', 'קרית אונו', 'הוד השרון', 'רמת גן', 'ר"ג', 'בני ברק', 'ב"ב', 'נווה ירק'
    ]
    for city in cities:
        name = re.sub(rf'\b{city}\b', '', name, flags=re.IGNORECASE)
    
    # הסרת סימני פיסוק מיותרים
    name = re.sub(r'[,\\/\-_]+', ' ', name)
    # הסרת רווחים מיותרים
    name = re.sub(r'\s+', ' ', name)
    name = name.strip()
    
    # הסרת מספרים וסימנים מיוחדים
    name = re.sub(r'[\d√\*\#\@]+', '', name)
    
    # אם נשארו פחות מ-2 תווים, החזר None
    if len(name) < 2:
        return None
    
    return name

def extract_contact_info(line):
    # מנסה לחלץ שם, טלפון, כתובת ומייל מהשורה
    parts = [part.strip() for part in line.split('|')]
    contact = {'שם': '', 'טלפון': '', 'כתובת': '', 'מייל': '', 'הערות': ''}
    
    # תבניות מורחבות לזיהוי כתובות מייל
    email_patterns = [
        # תבנית בסיסית
        r'[a-zA-Z0-9][a-zA-Z0-9._%+!#$&\'*+/=?^_`{|}~-]+@[a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?(?:\.[a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?)*\.[a-zA-Z]{2,}',
        
        # תבנית עם תוויות בעברית
        r'(?:מייל|אימייל|דוא"ל|דואל|mail|email|e-mail|כתובת מייל|כתובת אימייל)[\s:]*([a-zA-Z0-9][a-zA-Z0-9._%+!#$&\'*+/=?^_`{|}~-]+@[a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?(?:\.[a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?)*\.[a-zA-Z]{2,})',
        
        # תבנית עם תוויות באנגלית
        r'(?:E-?mail|Address|Contact)[\s:]*([a-zA-Z0-9][a-zA-Z0-9._%+!#$&\'*+/=?^_`{|}~-]+@[a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?(?:\.[a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?)*\.[a-zA-Z]{2,})',
        
        # תבנית עם מילות הקשר אחרי המייל
        r'([a-zA-Z0-9][a-zA-Z0-9._%+!#$&\'*+/=?^_`{|}~-]+@[a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?(?:\.[a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?)*\.[a-zA-Z]{2,})[\s:]*(?:בלבד|only|contact|קשר|ליצירת קשר|לפניות)'
    ]
    
    # חיפוש כתובות מייל
    for pattern in email_patterns:
        matches = re.finditer(pattern, line, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            email = match.group(1) if len(match.groups()) > 0 else match.group(0)
            email = email.lower().strip()
            
            # ניקוי תווים מיותרים
            email = re.sub(r'[,\s;]+$', '', email)  # הסרת תווי הפרדה בסוף
            email = re.sub(r'^[,\s;]+', '', email)  # הסרת תווי הפרדה בהתחלה
            
            if '@' in email and '.' in email.split('@')[1]:
                # בדיקות תקינות נוספות
                if len(email) > 254:  # RFC 5321
                    continue
                    
                try:
                    local, domain = email.split('@')
                    if not local or not domain:
                        continue
                        
                    if len(local) > 64:  # RFC 5321
                        continue
                        
                    if domain.startswith('.') or domain.endswith('.'):
                        continue
                        
                    if '..' in local or '..' in domain:
                        continue
                        
                    if local.startswith('.') or local.endswith('.'):
                        continue
                        
                    tld = domain.split('.')[-1]
                    if len(tld) < 2 or not tld.isalpha():
                        continue
                        
                    # הוספת המייל לרשימה
                    if contact['מייל']:
                        if email not in contact['מייל']:
                            contact['מייל'] += ';' + email
                    else:
                        contact['מייל'] = email
                        
                    # חיפוש שם בסביבת המייל
                    email_pos = line.find(email)
                    if email_pos != -1:
                        context_size = 100
                        before_email = line[max(0, email_pos - context_size):email_pos].strip()
                        after_email = line[email_pos + len(email):min(len(line), email_pos + len(email) + context_size)].strip()
                        
                        # חיפוש שם בעברית לפני המייל
                        name_match = re.search(r'([\u0590-\u05FF]+(?:\s+[\u0590-\u05FF]+){0,3})\s*$', before_email)
                        if name_match and not contact['שם']:
                            contact['שם'] = clean_name(name_match.group(1))
                        
                        # חיפוש שם בעברית אחרי המייל
                        if not contact['שם']:
                            name_match = re.search(r'^\s*([\u0590-\u05FF]+(?:\s+[\u0590-\u05FF]+){0,3})', after_email)
                            if name_match:
                                contact['שם'] = clean_name(name_match.group(1))
                        
                        # חיפוש שם בתוך תבניות נפוצות
                        if not contact['שם']:
                            name_patterns = [
                                r'שם:\s*([\u0590-\u05FF]+(?:\s+[\u0590-\u05FF]+){0,3})',
                                r'מאת:\s*([\u0590-\u05FF]+(?:\s+[\u0590-\u05FF]+){0,3})',
                                r'לכבוד:\s*([\u0590-\u05FF]+(?:\s+[\u0590-\u05FF]+){0,3})',
                                r'איש קשר:\s*([\u0590-\u05FF]+(?:\s+[\u0590-\u05FF]+){0,3})'
                            ]
                            for name_pattern in name_patterns:
                                name_match = re.search(name_pattern, before_email)
                                if name_match:
                                    contact['שם'] = clean_name(name_match.group(1))
                                    break
                                
                except Exception:
                    continue
                    
                line = line.replace(email, '').strip()
        if contact['מייל']:
            break
    
    # חיפוש מספר טלפון
    phone_patterns = [
        r'(?:050|052|054|057|03|04|077|076|072|073|074|075|058|051|053|055|056|02|08|09|1700|1800)-?\d{7}',  # תבנית רגילה
        r'05\d-?\d{7}',  # תבנית לנייד
        r'0\d-?\d{7}',   # תבנית לקווי
        r'0\d{2}-?\d{7}', # תבנית נוספת
        r'(?:טלפון|נייד|טל|פלאפון|טל׳|טל\'|Tel|Phone|Mobile)\s*:?\s*(0\d{1,2}[-.]?\d{7})',  # תבנית עם תווית
        r'(0\d{1,2}[-.]?\d{7})\s*(?:טלפון|נייד|טל|פלאפון|טל׳|טל\'|Tel|Phone|Mobile)'  # תבנית עם תווית אחרי
    ]
    
    for part in parts:
        # חיפוש מספר טלפון
        for pattern in phone_patterns:
            phone_match = re.search(pattern, part)
            if phone_match:
                phone = phone_match.group(1) if len(phone_match.groups()) > 0 else phone_match.group(0)
                # נרמול מספר הטלפון
                phone = re.sub(r'[-.]', '', phone)
                if len(phone) == 10:  # וידוא שזה מספר תקין
                    if contact['טלפון']:
                        if phone not in contact['טלפון']:
                            contact['טלפון'] += ';' + phone
                    else:
                        contact['טלפון'] = phone
                    # חיפוש שם בסביבת הטלפון
                    phone_pos = part.find(phone)
                    if phone_pos != -1:
                        context_size = 100
                        before_phone = part[max(0, phone_pos - context_size):phone_pos].strip()
                        after_phone = part[phone_pos + len(phone):min(len(part), phone_pos + len(phone) + context_size)].strip()
                        
                        # חיפוש שם בעברית לפני הטלפון
                        if not contact['שם']:
                            name_match = re.search(r'([\u0590-\u05FF]+(?:\s+[\u0590-\u05FF]+){0,3})\s*$', before_phone)
                            if name_match:
                                contact['שם'] = clean_name(name_match.group(1))
                        
                        # חיפוש שם בעברית אחרי הטלפון
                        if not contact['שם']:
                            name_match = re.search(r'^\s*([\u0590-\u05FF]+(?:\s+[\u0590-\u05FF]+){0,3})', after_phone)
                            if name_match:
                                contact['שם'] = clean_name(name_match.group(1))
                    
                    part = part.replace(phone, '').strip()
        
        # אם החלק מכיל מילים בעברית
        if re.search(r'[\u0590-\u05FF]', part):
            # חיפוש הערות
            notes_patterns = [
                r'(?:הערות|notes|note|הערה)[:]\s*(.*)',
                r'(?:לא מעוניין|לא רלוונטי|במעקב|בטיפול|נרשם|נרשמה)\s*(.*)'
            ]
            for pattern in notes_patterns:
                notes_match = re.search(pattern, part, re.IGNORECASE)
                if notes_match:
                    if contact['הערות']:
                        contact['הערות'] += '; ' + notes_match.group(1)
                    else:
                        contact['הערות'] = notes_match.group(1)
                    part = part.replace(notes_match.group(0), '').strip()
            
            if not contact['שם']:  # אם אין עדיין שם
                contact['שם'] = clean_name(part)
            elif not contact['כתובת']:  # אם אין כתובת
                contact['כתובת'] = part
    
    return contact

def extract_text_with_docx2txt(file_path):
    try:
        if file_path.endswith('.docx'):
            text = docx2txt.process(file_path)
            return text
        elif file_path.endswith('.doc'):
            # נסה להמיר את הקובץ ל-docx באמצעות LibreOffice
            docx_path = file_path + '.docx'
            try:
                subprocess.run(['soffice', '--headless', '--convert-to', 'docx', '--outdir', os.path.dirname(file_path), file_path], 
                             capture_output=True, check=True)
                if os.path.exists(docx_path):
                    text = docx2txt.process(docx_path)
                    os.remove(docx_path)  # מחיקת הקובץ הזמני
                    return text
            except Exception as e:
                print(f"Error converting DOC to DOCX: {str(e)}")
    except Exception as e:
        print(f"Error extracting text with docx2txt: {str(e)}")
    return None

def read_docx(file_path):
    try:
        # נסה קודם עם python-docx
        doc = Document(file_path)
        contacts = []
        full_text = ""
        
        # קריאת פסקאות
        for para in doc.paragraphs:
            if para.text.strip():
                full_text += para.text + "\n"
                contact = extract_contact_info(para.text)
                if contact['שם'] or contact['טלפון'] or contact['מייל']:
                    contacts.append(contact)
        
        # קריאת טבלאות
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text += row_text + "\n"
                    contact = extract_contact_info(row_text)
                    if contact['שם'] or contact['טלפון'] or contact['מייל']:
                        contacts.append(contact)
        
        # נסה גם עם docx2txt
        docx2txt_text = extract_text_with_docx2txt(file_path)
        if docx2txt_text:
            full_text += "\n" + docx2txt_text
        
        # חיפוש נוסף של כתובות מייל בכל הטקסט המלא
        email_patterns = [
            r'[a-zA-Z0-9][a-zA-Z0-9._%+-]*@(?:[a-zA-Z0-9-]+\.)+[a-zA-Z]{2,}',  # תבנית רגילה
            r'(?:מייל|אימייל|דוא"ל|דואל|mail|email|e-mail|כתובת מייל|כתובת אימייל)[\s:]*([a-zA-Z0-9][a-zA-Z0-9._%+-]*@(?:[a-zA-Z0-9-]+\.)+[a-zA-Z]{2,})',  # תבנית עם תווית
            r'([a-zA-Z0-9][a-zA-Z0-9._%+-]*@(?:[a-zA-Z0-9-]+\.)+[a-zA-Z]{2,})[\s:]*(?:מייל|אימייל|דוא"ל|דואל|mail|email|e-mail)',  # תבנית עם תווית אחרי
        ]
        
        emails_found = set()
        for pattern in email_patterns:
            matches = re.finditer(pattern, full_text, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                email = match.group(1) if len(match.groups()) > 0 else match.group(0)
                email = email.lower().strip()
                if '@' in email and '.' in email.split('@')[1]:
                    emails_found.add(email)
                    # חיפוש שם בסביבת כתובת המייל
                    context_size = 200
                    email_pos = full_text.find(email)
                    if email_pos != -1:
                        context = full_text[max(0, email_pos - context_size):min(len(full_text), email_pos + context_size)]
                        # חיפוש שם בעברית בהקשר
                        name_match = re.search(r'[\u0590-\u05FF]+(?:\s+[\u0590-\u05FF]+){1,3}', context)
                        if name_match:
                            name = name_match.group(0)
                            # בדיקה אם השם כבר קיים באנשי הקשר
                            found = False
                            for contact in contacts:
                                if contact['שם'] == name:
                                    contact['מייל'] = email
                                    found = True
                                    break
                            # אם לא נמצא איש קשר מתאים, נוסיף חדש
                            if not found:
                                contacts.append({'שם': name, 'טלפון': '', 'כתובת': '', 'מייל': email})
        
        return contacts
    except Exception as e:
        print(f"Error reading DOCX: {str(e)}")
        return []

def read_doc(file_path):
    try:
        # נסה קודם עם docx2txt
        full_text = extract_text_with_docx2txt(file_path)
        if not full_text:
            # אם docx2txt נכשל, נסה עם antiword
            result = subprocess.run(['antiword', file_path], capture_output=True, text=True)
            full_text = result.stdout.strip()
            
            if not full_text:
                # אם antiword נכשל, נסה עם catdoc
                result = subprocess.run(['catdoc', file_path], capture_output=True, text=True)
                full_text = result.stdout.strip()
        
        if full_text:
            contacts = []
            
            # עיבוד כל שורה
            for line in full_text.split('\n'):
                if line.strip():
                    contact = extract_contact_info(line)
                    if contact['שם'] or contact['טלפון'] or contact['מייל']:
                        contacts.append(contact)
            
            # חיפוש נוסף של כתובות מייל בכל הטקסט המלא
            email_patterns = [
                r'[a-zA-Z0-9][a-zA-Z0-9._%+-]*@(?:[a-zA-Z0-9-]+\.)+[a-zA-Z]{2,}',  # תבנית רגילה
                r'(?:מייל|אימייל|דוא"ל|דואל|mail|email|e-mail|כתובת מייל|כתובת אימייל)[\s:]*([a-zA-Z0-9][a-zA-Z0-9._%+-]*@(?:[a-zA-Z0-9-]+\.)+[a-zA-Z]{2,})',  # תבנית עם תווית
                r'([a-zA-Z0-9][a-zA-Z0-9._%+-]*@(?:[a-zA-Z0-9-]+\.)+[a-zA-Z]{2,})[\s:]*(?:מייל|אימייל|דוא"ל|דואל|mail|email|e-mail)',  # תבנית עם תווית אחרי
            ]
            
            emails_found = set()
            for pattern in email_patterns:
                matches = re.finditer(pattern, full_text, re.IGNORECASE | re.MULTILINE)
                for match in matches:
                    email = match.group(1) if len(match.groups()) > 0 else match.group(0)
                    email = email.lower().strip()
                    if '@' in email and '.' in email.split('@')[1]:
                        emails_found.add(email)
                        # חיפוש שם בסביבת כתובת המייל
                        context_size = 200
                        email_pos = full_text.find(email)
                        if email_pos != -1:
                            context = full_text[max(0, email_pos - context_size):min(len(full_text), email_pos + context_size)]
                            # חיפוש שם בעברית בהקשר
                            name_match = re.search(r'[\u0590-\u05FF]+(?:\s+[\u0590-\u05FF]+){1,3}', context)
                            if name_match:
                                name = name_match.group(0)
                                # בדיקה אם השם כבר קיים באנשי הקשר
                                found = False
                                for contact in contacts:
                                    if contact['שם'] == name:
                                        contact['מייל'] = email
                                        found = True
                                        break
                                # אם לא נמצא איש קשר מתאים, נוסיף חדש
                                if not found:
                                    contacts.append({'שם': name, 'טלפון': '', 'כתובת': '', 'מייל': email})
            
            return contacts
        
        return []
    except Exception as e:
        print(f"Error reading DOC: {str(e)}")
        return []

def format_excel(filename):
    # טעינת קובץ ה-Excel
    wb = load_workbook(filename)
    ws = wb.active
    
    # שינוי סדר העמודות והסרת עמודת הערות
    column_order = ['שם', 'טלפון', 'מייל', 'כתובת']
    current_columns = []
    for cell in ws[1]:
        current_columns.append(cell.value)
    
    # יצירת מיפוי עמודות
    column_mapping = {col: idx + 1 for idx, col in enumerate(current_columns)}
    
    # הוספת עמודות חדשות בסדר הרצוי
    for col_idx, col_name in enumerate(column_order, 1):
        if col_name not in current_columns:
            ws.insert_cols(col_idx)
            ws.cell(row=1, column=col_idx, value=col_name)
    
    # העתקת נתונים לסדר החדש
    data = []
    for row in ws.iter_rows(min_row=2):
        row_data = {}
        for cell in row:
            col_name = ws.cell(row=1, column=cell.column).value
            if col_name in column_order or col_name == 'מקור':  # כולל גם את עמודת המקור
                if col_name == 'מקור':
                    row_data['הערות'] = cell.value  # העברת תוכן המקור להערות
                else:
                    row_data[col_name] = cell.value
        data.append(row_data)
    
    # מחיקת כל השורות מלבד הכותרת
    ws.delete_rows(2, ws.max_row)
    
    # הוספת הנתונים בסדר החדש
    for row_idx, row_data in enumerate(data, 2):
        for col_idx, col_name in enumerate(column_order, 1):
            ws.cell(row=row_idx, column=col_idx, value=row_data.get(col_name, ''))
    
    # עיצוב כותרות העמודות
    header_font = Font(name='Arial', bold=True, size=12)
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    
    for col in range(1, ws.max_column + 1):
        cell = ws.cell(row=1, column=col)
        cell.font = Font(name='Arial', bold=True, size=12, color="FFFFFF")  # צבע לבן לטקסט
        cell.fill = header_fill
        cell.alignment = header_alignment
    
    # עיצוב תאי הטבלה
    data_font = Font(name='Arial', size=11)
    data_alignment = Alignment(horizontal="right", vertical="center", wrap_text=True)  # הוספת wrap_text=True
    alternate_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
    
    for row in range(2, ws.max_row + 1):
        for col in range(1, ws.max_column + 1):
            cell = ws.cell(row=row, column=col)
            cell.font = data_font
            cell.alignment = data_alignment
            # צביעת שורות לסירוגין
            if row % 2 == 0:
                cell.fill = alternate_fill
    
    # התאמת רוחב העמודות
    column_widths = {
        'שם': 25,
        'טלפון': 15,
        'מייל': 30,
        'כתובת': 20
    }
    
    for col_idx, col_name in enumerate(column_order, 1):
        column_letter = get_column_letter(col_idx)
        ws.column_dimensions[column_letter].width = column_widths[col_name]
    
    # הוספת מסגרת לטבלה
    thin_border = Border(
        left=Side(style='thin', color="BFBFBF"),
        right=Side(style='thin', color="BFBFBF"),
        top=Side(style='thin', color="BFBFBF"),
        bottom=Side(style='thin', color="BFBFBF")
    )
    
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
        for cell in row:
            cell.border = thin_border
    
    # הקפאת שורת הכותרת
    ws.freeze_panes = 'A2'
    
    # שמירת השינויים
    wb.save(filename)

def compare_with_original(df, examples_dir):
    print("\nComparing with original files:")
    print("-" * 50)
    
    # קריאת כל הקבצים המקוריים
    all_original_text = ""
    for filename in os.listdir(examples_dir):
        if filename.endswith('.docx') or filename.endswith('.doc'):
            file_path = os.path.join(examples_dir, filename)
            print(f"\nAnalyzing file: {filename}")
            
            # קריאת הטקסט מהקובץ
            if filename.endswith('.docx'):
                text = docx2txt.process(file_path)
            else:
                text = extract_text_with_docx2txt(file_path)
            
            if text:
                all_original_text += text + "\n"
                
                # חיפוש טלפונים בקובץ המקורי
                phone_numbers = re.findall(r'(?:050|052|054|057|03|04|077|076|072|073|074|075|058|051|053|055|056|02|08|09|1700|1800)-?\d{7}|05\d-?\d{7}', text)
                phones_in_excel = df[df['טלפון'].notna()]['טלפון'].tolist()
                
                # חיפוש כתובות מייל בקובץ המקורי
                emails = re.findall(r'[a-zA-Z0-9][a-zA-Z0-9._%+-]*@(?:[a-zA-Z0-9-]+\.)+[a-zA-Z]{2,}', text)
                emails_in_excel = df[df['מייל'].notna()]['מייל'].tolist()
                
                # הדפסת סטטיסטיקות לקובץ
                print(f"Found in original file:")
                print(f"- Phone numbers: {len(phone_numbers)}")
                print(f"- Email addresses: {len(emails)}")
                
                # בדיקת מספרי טלפון שלא נכנסו לטבלה
                missing_phones = [phone for phone in phone_numbers if phone not in phones_in_excel]
                if missing_phones:
                    print("\nPhone numbers found in file but missing from Excel:")
                    for phone in missing_phones[:5]:  # מציג רק 5 דוגמאות
                        context = get_context(text, phone)
                        print(f"- {phone} (Context: {context})")
                
                # בדיקת כתובות מייל שלא נכנסו לטבלה
                missing_emails = [email.lower() for email in emails if email.lower() not in [e.lower() for e in emails_in_excel]]
                if missing_emails:
                    print("\nEmail addresses found in file but missing from Excel:")
                    for email in missing_emails[:5]:  # מציג רק 5 דוגמאות
                        context = get_context(text, email)
                        print(f"- {email} (Context: {context})")

def get_context(text, search_term, context_size=50):
    """מחזיר את ההקשר סביב מונח החיפוש"""
    pos = text.find(search_term)
    if pos != -1:
        start = max(0, pos - context_size)
        end = min(len(text), pos + len(search_term) + context_size)
        context = text[start:end].replace('\n', ' ').strip()
        return f"...{context}..."
    return ""

def main():
    examples_dir = "../מתווכים"
    all_contacts = []
    
    for filename in os.listdir(examples_dir):
        if filename.endswith('.docx') or filename.endswith('.doc'):
            file_path = os.path.join(examples_dir, filename)
            print(f"\nReading file: {filename}")
            
            if filename.endswith('.docx'):
                contacts = read_docx(file_path)
            else:
                contacts = read_doc(file_path)
            
            # הוספת שם הקובץ כמקור
            for contact in contacts:
                contact['הערות'] = filename  # שינוי מ'מקור' ל'הערות'
            
            all_contacts.extend(contacts)
    
    # יצירת DataFrame
    df = pd.DataFrame(all_contacts)
    
    # ניקוי נוסף של השמות
    df['שם'] = df['שם'].apply(clean_name)
    
    # הסרת שורות ריקות או חלקיות
    df = df.dropna(subset=['שם'])
    
    # הסרת אנשי קשר שאין להם טלפון או מייל
    df = df[df['טלפון'].fillna('').str.len().gt(0) | df['מייל'].fillna('').str.len().gt(0)]
    
    # איחוד רשומות כפולות
    df = df.groupby('שם', as_index=False).agg({
        'טלפון': lambda x: ';'.join(filter(None, set(x.fillna('').str.split(';').explode()))),
        'מייל': lambda x: ';'.join(filter(None, set(x.fillna('').str.split(';').explode()))),
        'כתובת': lambda x: ';'.join(filter(None, set(x.fillna('')))),
        'הערות': lambda x: ';'.join(set(x.fillna('')))
    })
    
    # מיון לפי שם
    df = df.sort_values('שם')
    
    # סידור סדר העמודות
    df = df[['שם', 'טלפון', 'מייל', 'כתובת']]
    
    # שמירה לקובץ Excel
    output_file = 'contact_results.xlsx'
    df.to_excel(output_file, index=False)
    
    # עיצוב הקובץ
    format_excel(output_file)
    
    print(f"\nSaved {len(df)} contacts to {output_file}")
    
    # השוואה עם הקבצים המקוריים
    compare_with_original(df, examples_dir)

if __name__ == "__main__":
    main() 